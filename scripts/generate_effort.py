from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
sec = doc.sections[0]
sec.top_margin    = Cm(1.8)
sec.bottom_margin = Cm(1.8)
sec.left_margin   = Cm(2.0)
sec.right_margin  = Cm(2.0)

# ── Helpers ───────────────────────────────────────────────────────────────────
def cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

def cell_border(cell, color='BFBFBF', sz='4'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0'); b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def set_col_width(table, col_widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(col_widths_cm):
                cell.width = Cm(col_widths_cm[i])

def add_heading(doc, text, level=1, color="1F497D"):
    p = doc.add_heading(text, level)
    p.runs[0].font.color.rgb = RGBColor(
        int(color[0:2],16), int(color[2:4],16), int(color[4:6],16))
    return p

def para(doc, text, bold=False, italic=False, fs=10, color="333333", space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(fs)
    r.font.color.rgb = RGBColor(int(color[0:2],16), int(color[2:4],16), int(color[4:6],16))
    return p

def bullet(doc, text, fs=9):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(fs)
    return p

# ─── Role colours ──────────────────────────────────────────────────────────────
R = {
    'SA':   ('1F497D', 'D6E4F7'),
    'AIL':  ('6A1B9A', 'EDE7F6'),
    'AIE1': ('00695C', 'E0F2F1'),
    'AIE2': ('2E7D32', 'E8F5E9'),
    'QA':   ('C62828', 'FFEBEE'),
}

PHASE_COLOR  = '2E4057'
HEADER_COLOR = '1F497D'
TOTAL_COLOR  = 'BF360C'
ALT_ROW      = 'F5F9FF'

# =============================================================================
# PAGE 1 — TITLE + TEAM
# =============================================================================
title = doc.add_heading('SHARP EV AI Chatbot', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_heading('Project Effort Estimation', 2)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(
    'Prepared by: Solution Architect   |   Date: 2026-06-22   |   '
    'Version: 1.0   |   Unit: Person-Weeks (PW)  ·  1 PW = 5 Person-Days')
r.font.size = Pt(8.5); r.font.italic = True
r.font.color.rgb = RGBColor(0x70,0x70,0x70)
doc.add_paragraph()

# ── Team roster box ───────────────────────────────────────────────────────────
team_tbl = doc.add_table(rows=1, cols=5)
team_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
team_tbl.style = 'Table Grid'
set_col_width(team_tbl, [3.3]*5)

roles = [
    ('SA', 'Solution Architect',    '1 person',  '1F497D'),
    ('AIL','AI Lead',               '1 person',  '6A1B9A'),
    ('AIE','AI Engineer',           '2 persons', '00695C'),
    ('QA', 'QA Engineer',           '1 person',  'C62828'),
    ('',   'Total Team',            '5 persons', '37474F'),
]
for i, (abbr, role, count, col) in enumerate(roles):
    c = team_tbl.rows[0].cells[i]
    cell_bg(c, col); cell_border(c, col)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{abbr}\n'); r.bold=True; r.font.size=Pt(13)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    r2 = p.add_run(f'{role}\n'); r2.font.size=Pt(8)
    r2.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    r3 = p.add_run(count); r3.font.size=Pt(9); r3.bold=True
    r3.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

doc.add_paragraph()

# =============================================================================
# SECTION 1 — EFFORT SUMMARY BY PHASE
# =============================================================================
add_heading(doc, '1.  Effort Summary by Phase', 1)

# Columns: Phase | Weeks | SA | AIL | AIE1 | AIE2 | QA | Total PW
phases = [
    # (phase_name, cal_weeks, SA, AIL, AIE1, AIE2, QA)
    # Totals: SA=9.0 | AIL=17.0 | AIE1=15.0 | AIE2=13.0 | QA=16.0 = 70.0 PW
    ('Phase 1 — Discovery & Architecture',      2,   2.0,  2.0,  1.0,  1.0,  1.0),
    ('Phase 2 — Infrastructure & Env Setup',    2,   2.0,  1.0,  1.5,  1.5,  1.0),
    ('Phase 3A — Self-Diagnosis Module',         4,   0.5,  3.0,  3.0,  1.5,  2.0),
    ('Phase 3B — EV Usage Consultation (RAG)',   5,   0.5,  4.5,  3.5,  2.0,  2.5),
    ('Phase 3C — General Chat',                  2,   0.0,  1.0,  0.5,  1.5,  0.5),
    ('Phase 3D — General Automotive Guidance',   2,   0.0,  1.0,  1.5,  0.5,  1.0),
    ('Phase 4 — Voice I/O & Security',           2,   1.5,  1.0,  1.5,  1.5,  1.0),
    ('Phase 5 — Integration Testing & UAT',      3,   1.0,  2.5,  2.0,  2.0,  4.5),
    ('Phase 6 — Production Deployment',          2,   1.5,  1.0,  0.5,  1.5,  2.5),
]

COL_W = [6.0, 1.4, 1.4, 1.4, 1.4, 1.4, 1.4, 1.5]
headers = ['Phase', 'Cal.\nWeeks', 'SA\n(PW)', 'AI Lead\n(PW)', 'AIE 1\n(PW)', 'AIE 2\n(PW)', 'QA\n(PW)', 'Phase\nTotal']

tbl = doc.add_table(rows=1+len(phases)+1, cols=8)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_col_width(tbl, COL_W)

# Header row
hrow = tbl.rows[0]
for i, h in enumerate(headers):
    c = hrow.cells[i]; cell_bg(c, PHASE_COLOR); cell_border(c, PHASE_COLOR)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.bold=True; r.font.size=Pt(8.5)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

# Data rows
tot_sa=tot_ail=tot_aie1=tot_aie2=tot_qa=tot_cal=0
for ri, (pname, cal, sa, ail, aie1, aie2, qa) in enumerate(phases):
    row = tbl.rows[ri+1]
    phase_total = sa+ail+aie1+aie2+qa
    tot_sa+=sa; tot_ail+=ail; tot_aie1+=aie1; tot_aie2+=aie2; tot_qa+=qa; tot_cal+=cal
    bg = 'FFFFFF' if ri%2==0 else ALT_ROW
    vals = [pname, str(cal), str(sa), str(ail), str(aie1), str(aie2), str(qa), str(phase_total)]
    role_cols = {2: R['SA'][1], 3: R['AIL'][1], 4: R['AIE1'][1], 5: R['AIE2'][1], 6: R['QA'][1]}
    for ci, v in enumerate(vals):
        c = row.cells[ci]
        cell_bg(c, role_cols.get(ci, bg)); cell_border(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.CENTER
        is_bold = (ci==0 or ci==7)
        col_hex = (R['SA'][0] if ci==2 else R['AIL'][0] if ci==3 else
                   R['AIE1'][0] if ci==4 else R['AIE2'][0] if ci==5 else
                   R['QA'][0] if ci==6 else TOTAL_COLOR if ci==7 else '333333')
        r2 = p.add_run(v); r2.font.size=Pt(9); r2.bold=is_bold
        r2.font.color.rgb = RGBColor(int(col_hex[0:2],16),int(col_hex[2:4],16),int(col_hex[4:6],16))

# Totals row
grand_total = tot_sa+tot_ail+tot_aie1+tot_aie2+tot_qa
tot_row = tbl.rows[-1]
tot_vals = ['TOTAL', str(tot_cal), str(tot_sa), str(tot_ail), str(tot_aie1), str(tot_aie2), str(tot_qa), str(grand_total)]
tot_labels = ['TOTAL PERSON-WEEKS', f'{tot_cal} wks',
              f'{tot_sa} PW\n({int(tot_sa*5)} days)',
              f'{tot_ail} PW\n({int(tot_ail*5)} days)',
              f'{tot_aie1} PW\n({int(tot_aie1*5)} days)',
              f'{tot_aie2} PW\n({int(tot_aie2*5)} days)',
              f'{tot_qa} PW\n({int(tot_qa*5)} days)',
              f'{grand_total} PW\n({int(grand_total*5)} days)']
role_bg = {2: R['SA'][0], 3: R['AIL'][0], 4: R['AIE1'][0], 5: R['AIE2'][0], 6: R['QA'][0]}
for ci, lbl in enumerate(tot_labels):
    c = tot_row.cells[ci]
    cell_bg(c, role_bg.get(ci, TOTAL_COLOR)); cell_border(c, TOTAL_COLOR, '6')
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.CENTER
    r2 = p.add_run(lbl); r2.bold=True; r2.font.size=Pt(9)
    r2.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

doc.add_paragraph()

# =============================================================================
# SECTION 2 — TIMELINE (GANTT SUMMARY)
# =============================================================================
add_heading(doc, '2.  Project Timeline  (16 Weeks / ~4 Months)', 1)

WEEKS = 16
gantt_data = [
    # (label, start_week, end_week, color)
    ('Phase 1: Discovery & Architecture',   1,  2,  '1F497D'),
    ('Phase 2: Infrastructure Setup',        3,  4,  '6A1B9A'),
    ('Phase 3A: Self-Diagnosis Module',      5,  8,  'C62828'),
    ('Phase 3B: EV Usage Consult (RAG)',     5,  9,  '1565C0'),
    ('Phase 3C: General Chat',               6,  7,  '2E7D32'),
    ('Phase 3D: General Auto Guidance',      7,  8,  'E65100'),
    ('Phase 4: Voice I/O & Security',        8,  9,  '37474F'),
    ('Phase 5: Integration Testing & UAT',  10, 12,  '880E4F'),
    ('Phase 6: Production Deployment',      13, 16,  'BF360C'),
]

LABEL_W = Cm(5.2)
CELL_W  = Cm(0.65)
gantt_cols = 1 + WEEKS
gtbl = doc.add_table(rows=1+len(gantt_data), cols=gantt_cols)
gtbl.style = 'Table Grid'
gtbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
gh = gtbl.rows[0]
gh.cells[0].width = LABEL_W
cell_bg(gh.cells[0], PHASE_COLOR); cell_border(gh.cells[0], PHASE_COLOR)
p = gh.cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Phase / Activity'); r.bold=True; r.font.size=Pt(8)
r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
for w in range(WEEKS):
    c = gh.cells[w+1]; c.width = CELL_W
    cell_bg(c, PHASE_COLOR); cell_border(c, PHASE_COLOR)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'W{w+1}'); r.bold=True; r.font.size=Pt(7)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

# Gantt rows
for ri, (label, s, e, col) in enumerate(gantt_data):
    row = gtbl.rows[ri+1]
    row.cells[0].width = LABEL_W
    cell_bg(row.cells[0], 'F5F5F5'); cell_border(row.cells[0])
    p = row.cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(label); r.font.size=Pt(8); r.bold=True
    r.font.color.rgb = RGBColor(int(col[0:2],16),int(col[2:4],16),int(col[4:6],16))
    for w in range(WEEKS):
        c = row.cells[w+1]; c.width = CELL_W
        if s <= w+1 <= e:
            cell_bg(c, col); cell_border(c, col)
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p.add_run('▮'); r2.font.size=Pt(7)
            r2.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        else:
            cell_bg(c, 'F5F5F5'); cell_border(c, 'DDDDDD')

doc.add_paragraph()

# =============================================================================
# SECTION 3 — PHASE-BY-PHASE ACTIVITY BREAKDOWN
# =============================================================================
add_heading(doc, '3.  Phase-by-Phase Activity Breakdown', 1)

phase_details = [
    ('Phase 1 — Discovery & Architecture', '1F497D', 'Week 1–2', [
        ('SA (2 PW)',   ['Requirements workshops with SHARP client',
                         'AWS architecture design (HLD + LLD)',
                         'Multi-modal RAG design review',
                         'Cloud account setup + IAM strategy']),
        ('AI Lead (2 PW)', ['AI/ML stack selection (Bedrock, SageMaker, Kendra)',
                             'Multi-modal RAG architecture design',
                             'Embedding model selection (Titan v2 + Multimodal)',
                             'LLM evaluation framework design']),
        ('AIE 1 (1 PW)', ['Research spike: Bedrock + RAG patterns',
                           'Vehicle telemetry API feasibility check']),
        ('AIE 2 (1 PW)', ['Research spike: COCORO MEMBERS ID OAuth flow',
                           'API Gateway + Lambda architecture spike']),
        ('QA (1 PW)',    ['Test strategy document',
                           'Test environment requirements',
                           'Acceptance criteria definition']),
    ]),
    ('Phase 2 — Infrastructure & Environment Setup', '6A1B9A', 'Week 3–4', [
        ('SA (2 PW)',   ['VPC, subnets, security groups, IAM roles',
                         'CI/CD pipeline (CodePipeline + CodeBuild)',
                         'CloudWatch dashboards + alerting baseline',
                         'WAF + Shield baseline setup']),
        ('AI Lead (1 PW)', ['Amazon Bedrock model access + quota requests',
                             'OpenSearch Serverless cluster provisioning',
                             'SageMaker Studio workspace setup']),
        ('AIE 1 (1.5 PW)', ['AWS IoT Core + Kinesis stream setup',
                             'Lambda project scaffold + deployment framework',
                             'S3 buckets + lifecycle policies']),
        ('AIE 2 (1.5 PW)', ['API Gateway + Lambda scaffold',
                             'Amazon Cognito + COCORO MEMBERS ID integration',
                             'DynamoDB tables (sessions, conversation history)']),
        ('QA (1 PW)',  ['Test environment provisioned (dev / staging)',
                         'Test data preparation (mock vehicle telemetry)',
                         'Postman / test automation framework setup']),
    ]),
    ('Phase 3A — Self-Diagnosis Module', 'C62828', 'Week 5–8', [
        ('SA (0.5 PW)',  ['Architecture review — vehicle data flow',
                           'Security sign-off for vehicle telemetry pipeline']),
        ('AI Lead (3 PW)', ['SageMaker fault diagnosis ML model training + tuning',
                             'Amazon Rekognition integration (warning light image analysis)',
                             'Diagnosis confidence scoring logic',
                             'Severity classification (Critical / Warning / Info)']),
        ('AIE 1 (3 PW)', ['IoT Core → Kinesis → Lambda pipeline build',
                           'Fault code database (Aurora) schema + ingestion',
                           'Step Functions diagnosis workflow orchestration',
                           'Vehicle image upload (S3) + Rekognition trigger']),
        ('AIE 2 (1.5 PW)', ['Diagnosis API (REST) — request/response contract',
                             'Dealer booking integration stub',
                             'Diagnosis history storage (DynamoDB + Aurora)']),
        ('QA (2 PW)',    ['Test cases: 50+ fault scenarios (OBD-II codes)',
                           'Image recognition accuracy tests (warning lights)',
                           'End-to-end diagnosis flow testing',
                           'Severity classification validation']),
    ]),
    ('Phase 3B — EV Usage Consultation (Multi-modal RAG)', '1565C0', 'Week 5–9', [
        ('SA (0.5 PW)',  ['Data residency review — manual storage compliance',
                           'RAG architecture sign-off']),
        ('AI Lead (4.5 PW)', ['Document ingestion pipeline design + orchestration',
                               'Amazon Textract — PDF OCR + table extraction',
                               'Titan Text Embeddings v2 — chunk generation',
                               'Titan Multimodal Embeddings — vehicle diagram indexing',
                               'Amazon Kendra index setup + manual ingestion',
                               'Cohere Rerank v3 integration for re-scoring',
                               'Prompt augmentation templates per module',
                               'RAG accuracy tuning (recall / precision)']),
        ('AIE 1 (3.5 PW)', ['Document loader Lambda (PDF → chunks → embeddings)',
                             'OpenSearch kNN index build + mapping configuration',
                             'Chunk metadata store (DynamoDB)',
                             'Index refresh pipeline (new manual versions)']),
        ('AIE 2 (2 PW)', ['Retrieval API — query → embed → kNN → rerank',
                           'Response formatter — answer + page citation + URL',
                           'Fallback handler (no-answer / low-confidence)']),
        ('QA (2.5 PW)',  ['RAG accuracy benchmark testing (over XX% target)',
                           'Citation accuracy — correct page references',
                           'Multi-model (text + image query) testing',
                           'Fallback response quality validation']),
    ]),
    ('Phase 3C — General Chat', '2E7D32', 'Week 6–7', [
        ('AI Lead (1 PW)', ['Bedrock Claude 3 system prompt design',
                             'Guardrails configuration (topic scoping)',
                             'Content moderation rules (METI compliance)']),
        ('AIE 1 (0.5 PW)', ['Chat session API integration support']),
        ('AIE 2 (1.5 PW)', ['Chat session API — WebSocket + REST',
                             'Multi-turn conversation history (DynamoDB)',
                             'Usage limit enforcement (XX req/month quota)']),
        ('QA (0.5 PW)',   ['Conversation flow testing (multi-turn)',
                           'Guardrail effectiveness testing']),
    ]),
    ('Phase 3D — General Automotive Guidance', 'E65100', 'Week 7–8', [
        ('AI Lead (1 PW)', ['Automotive knowledge base ingestion + indexing',
                             'Domain-specific RAG prompt tuning',
                             'Cohere multilingual embed for JP/EN content']),
        ('AIE 1 (1.5 PW)', ['KB document loader + chunking pipeline',
                             'OpenSearch automotive index + metadata schema']),
        ('AIE 2 (0.5 PW)', ['Automotive guidance API + response formatter']),
        ('QA (1 PW)',     ['Automotive guidance accuracy testing',
                           'Cross-module boundary testing (chat vs auto)']),
    ]),
    ('Phase 4 — Voice I/O & Security', '37474F', 'Week 8–9', [
        ('SA (1.5 PW)',   ['AWS WAF rules (geo-block, rate limit, SQLi)',
                           'AWS Shield Advanced activation',
                           'Amazon GuardDuty threat detection config',
                           'Security architecture review + pen test scope']),
        ('AI Lead (1 PW)', ['Amazon Transcribe — Japanese STT model tuning',
                             'Amazon Polly — TTS voice character selection + testing']),
        ('AIE 1 (1.5 PW)', ['Transcribe streaming API integration',
                             'Voice session auto-deactivation timer (TBD value)',
                             'STT accuracy validation pipeline']),
        ('AIE 2 (1.5 PW)', ['Polly TTS integration — text → voice response',
                             'Voice + text parallel response delivery',
                             'Rate limiter Lambda (XX req/month enforcement)']),
        ('QA (1 PW)',     ['STT accuracy testing (over XX% target)',
                           'TTS quality review (voice + transcript sync)',
                           'Security scan (OWASP, replay attack tests)']),
    ]),
    ('Phase 5 — Integration Testing & UAT', '880E4F', 'Week 10–12', [
        ('SA (1 PW)',     ['End-to-end architecture validation',
                           'Performance & load test review',
                           'UAT environment sign-off']),
        ('AI Lead (2.5 PW)', ['Full RAG pipeline accuracy benchmark (all 4 modules)',
                               'LLM response quality review + prompt tuning',
                               'UAT support — AI-related defect triage']),
        ('AIE 1 (2 PW)',  ['End-to-end integration — cross-module testing',
                           'Bug fixes — ingestion + retrieval pipeline',
                           'Performance profiling + bottleneck resolution']),
        ('AIE 2 (2 PW)',  ['Bug fixes — API, auth, voice, usage limits',
                           'Cross-browser / cross-device compatibility',
                           'Load testing support']),
        ('QA (4.5 PW)',   ['Full regression test suite execution',
                           'Performance testing (concurrent users load test)',
                           'Security penetration testing',
                           'UAT facilitation with SHARP stakeholders',
                           'Defect management + retesting']),
    ]),
    ('Phase 6 — Production Deployment & Handover', 'BF360C', 'Week 13–16', [
        ('SA (1.5 PW)',   ['Production deployment runbook',
                           'CloudWatch production dashboards + alarms',
                           'Cost optimisation review (LLM API + infra)',
                           'Architecture handover documentation']),
        ('AI Lead (1 PW)', ['Production model deployment + smoke test',
                             'RAG accuracy production validation',
                             'AI runbook + model retraining guide']),
        ('AIE 1 (0.5 PW)', ['Production deployment + rollback support']),
        ('AIE 2 (1.5 PW)', ['Production deployment + smoke testing',
                             'Monitoring dashboards handover']),
        ('QA (2.5 PW)',   ['Production smoke + sanity testing',
                           'Go/No-Go sign-off checklist',
                           'Test evidence pack for METI compliance']),
    ]),
]

for (phase_name, col, timeline, role_activities) in phase_details:
    # Phase header
    ph_tbl = doc.add_table(rows=1, cols=1); ph_tbl.style = 'Table Grid'
    ph_c = ph_tbl.rows[0].cells[0]
    cell_bg(ph_c, col); cell_border(ph_c, col)
    p = ph_c.paragraphs[0]
    r = p.add_run(f'  {phase_name}  ·  {timeline}')
    r.bold=True; r.font.size=Pt(10.5)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

    for role_label, activities in role_activities:
        # Role sub-header
        rh_tbl = doc.add_table(rows=1, cols=1); rh_tbl.style = 'Table Grid'
        rh_c = rh_tbl.rows[0].cells[0]
        rkey = role_label.split(' ')[0]  # SA, AI, AIE, QA
        bg = (R['SA'][1] if rkey=='SA' else R['AIL'][1] if 'Lead' in role_label
              else R['AIE1'][1] if '1' in role_label else R['AIE2'][1] if '2' in role_label
              else R['QA'][1])
        fc = (R['SA'][0] if rkey=='SA' else R['AIL'][0] if 'Lead' in role_label
              else R['AIE1'][0] if '1' in role_label else R['AIE2'][0] if '2' in role_label
              else R['QA'][0])
        cell_bg(rh_c, bg); cell_border(rh_c, fc)
        p = rh_c.paragraphs[0]
        r = p.add_run(f'  {role_label}')
        r.bold=True; r.font.size=Pt(9)
        r.font.color.rgb = RGBColor(int(fc[0:2],16),int(fc[2:4],16),int(fc[4:6],16))

        # Activity table
        act_tbl = doc.add_table(rows=len(activities), cols=1)
        act_tbl.style = 'Table Grid'
        for ai, act in enumerate(activities):
            c = act_tbl.rows[ai].cells[0]
            cell_bg(c, 'FFFFFF' if ai%2==0 else 'F9F9F9')
            cell_border(c, 'DDDDDD', '2')
            p = c.paragraphs[0]
            r = p.add_run(f'  ▸  {act}')
            r.font.size=Pt(9)

    doc.add_paragraph()

# =============================================================================
# SECTION 4 — TOTAL EFFORT SUMMARY
# =============================================================================
add_heading(doc, '4.  Total Effort Summary', 1)

summary_data = [
    ('Solution Architect',  '1', str(tot_sa), str(int(tot_sa*5)), str(int(tot_sa*5*1)), 'Part-time from Phase 3 onwards'),
    ('AI Lead',             '1', str(tot_ail), str(int(tot_ail*5)), str(int(tot_ail*5*1)), 'Full-time — critical path owner'),
    ('AI Engineer 1',       '1', str(tot_aie1),str(int(tot_aie1*5)),'N/A', 'Full-time — pipeline & IoT specialist'),
    ('AI Engineer 2',       '1', str(tot_aie2),str(int(tot_aie2*5)),'N/A', 'Full-time — API & auth specialist'),
    ('QA Engineer',         '1', str(tot_qa), str(int(tot_qa*5)), 'N/A', 'Full-time from Phase 2 onwards'),
]
grand_pw = tot_sa + tot_ail + tot_aie1 + tot_aie2 + tot_qa
grand_pd = int(grand_pw * 5)

sum_tbl = doc.add_table(rows=1+len(summary_data)+1, cols=6)
sum_tbl.style = 'Table Grid'
sum_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_col_width(sum_tbl, [4.0, 1.2, 1.5, 1.5, 1.8, 5.0])

sum_hdrs = ['Role', 'Count', 'Total PW', 'Total Days', 'Availability', 'Notes']
shr = sum_tbl.rows[0]
for i, h in enumerate(sum_hdrs):
    c = shr.cells[i]; cell_bg(c, PHASE_COLOR); cell_border(c)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.bold=True; r.font.size=Pt(9)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

role_colors = [R['SA'][1], R['AIL'][1], R['AIE1'][1], R['AIE2'][1], R['QA'][1]]
role_fcolors= [R['SA'][0], R['AIL'][0], R['AIE1'][0], R['AIE2'][0], R['QA'][0]]

for ri, (role, cnt, pw, pd, avail, note) in enumerate(summary_data):
    row = sum_tbl.rows[ri+1]
    vals = [role, cnt, f'{pw} PW', f'{pd} days', avail, note]
    for ci, v in enumerate(vals):
        c = row.cells[ci]; cell_bg(c, role_colors[ri]); cell_border(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci in [0,5] else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(v); r.font.size=Pt(9)
        r.bold = (ci in [0,2,3])
        r.font.color.rgb = RGBColor(int(role_fcolors[ri][0:2],16),
                                    int(role_fcolors[ri][2:4],16),
                                    int(role_fcolors[ri][4:6],16))

tot_r = sum_tbl.rows[-1]
tot_entries = ['GRAND TOTAL', '5', f'{grand_pw} PW', f'{grand_pd} days', '16 Weeks', 'Project Duration: ~4 Months']
for ci, v in enumerate(tot_entries):
    c = tot_r.cells[ci]; cell_bg(c, TOTAL_COLOR); cell_border(c, TOTAL_COLOR, '6')
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci in [0,5] else WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(v); r.bold=True; r.font.size=Pt(9.5)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

doc.add_paragraph()

# =============================================================================
# SECTION 5 — KEY ASSUMPTIONS
# =============================================================================
add_heading(doc, '5.  Key Assumptions', 1)

assumptions = [
    ('Scope',          ['All 4 modules (Self-diagnosis, EV Usage Consultation, Chat, General Auto Guidance) are in scope for v1.',
                         'Voice I/O (Japanese STT + TTS) is included in v1.',
                         'COCORO MEMBERS ID auth integration is in scope; SSO API spec provided by SHARP.']),
    ('Client Inputs',  ['SHARP provides EV manuals in digital format (PDF/HTML) within Week 1.',
                         'Vehicle telemetry API spec / sandbox provided by Week 2.',
                         'All TBD values (accuracy %, response time, concurrent users, usage cap) confirmed in Phase 1.',
                         'SHARP assigns a technical point-of-contact for weekly progress reviews.']),
    ('Technology',     ['Amazon Bedrock (Claude 3 Sonnet/Haiku) access approved and quota granted before Phase 3.',
                         'No custom LLM fine-tuning required — prompt engineering and RAG are sufficient.',
                         'AWS region: ap-northeast-1 (Tokyo) for data residency compliance.',
                         'Existing AWS account provided by SHARP or new account created in Phase 1.']),
    ('Team',           ['All 5 team members are available full-time from project start.',
                         '2 AI Engineers have prior experience with AWS Lambda, API Gateway, and Python.',
                         'AI Lead has hands-on experience with RAG pipelines and Amazon Bedrock.',
                         'No team ramp-up time required beyond the Phase 1 research spike.']),
    ('Excluded',       ['Mobile app (iOS/Android) UI development — assumed to be handled by SHARP or a separate team.',
                         'In-car head unit integration — out of scope for v1.',
                         'Custom LLM training or fine-tuning.',
                         'Dealer booking system integration (stub API only in v1).',
                         'Multi-region deployment — Japan (ap-northeast-1) only for v1.']),
]

for category, points in assumptions:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f'  {category}')
    r.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    for point in points:
        bullet(doc, point)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# =============================================================================
# SECTION 6 — KEY RISKS
# =============================================================================
add_heading(doc, '6.  Key Risks', 1)

risks = [
    ('HIGH',   'TBD values not confirmed by SHARP in Phase 1 (accuracy %, response time, concurrent users).', 'Resolve all XX placeholders in Phase 1 workshops. Block Phase 3 start if unresolved.'),
    ('HIGH',   'Vehicle telemetry API does not exist — needs to be built by SHARP.', 'Phase 3A (Self-diagnosis) is gated on this API. Add 3–4 PW buffer if SHARP needs to build it.'),
    ('HIGH',   'EV manuals provided as scanned PDFs — OCR quality degrades RAG accuracy.', 'Demand machine-readable source in Phase 1. If scanned only, add 2 PW for Textract QA.'),
    ('MEDIUM', 'Amazon Bedrock Claude 3 quota limits delay LLM integration.', 'Request quota increase in Week 1. Maintain fallback to Claude 2 or GPT-4 via API Gateway.'),
    ('MEDIUM', 'COCORO MEMBERS ID OAuth spec delayed or incomplete.', 'Use mock auth in dev environment. Block production UAT until full spec received.'),
    ('MEDIUM', 'Japanese STT (Transcribe) accuracy below target for automotive vocabulary.', 'Add custom vocabulary list in Phase 4. Budget 1 additional PW for AIL tuning.'),
    ('LOW',    'OpenSearch kNN search latency exceeds response time target.', 'Add ElastiCache semantic cache layer (planned). Optimise top-K value and shard count.'),
    ('LOW',    'METI AI compliance requirements more stringent than expected.', 'Schedule legal review in Phase 1. Add 1 PW SA + 1 PW QA buffer for compliance evidence.'),
]

risk_tbl = doc.add_table(rows=1+len(risks), cols=4)
risk_tbl.style = 'Table Grid'
risk_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_col_width(risk_tbl, [1.5, 5.5, 6.0, 3.5])

for i, h in enumerate(['Level', 'Risk', 'Mitigation', 'Impacted Phase']):
    c = risk_tbl.rows[0].cells[i]; cell_bg(c, PHASE_COLOR); cell_border(c)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.bold=True; r.font.size=Pt(9)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

risk_impacts = ['Phase 1', 'Phase 3A', 'Phase 3B', 'Phase 3 start', 'Phase 4–5', 'Phase 4', 'Phase 5', 'Phase 1']
risk_fill = {'HIGH': ('FFCDD2','C62828'), 'MEDIUM': ('FFF9C4','E65100'), 'LOW': ('E8F5E9','2E7D32')}
for ri, (level, risk, mitigation) in enumerate(risks):
    row = risk_tbl.rows[ri+1]
    bg, fc = risk_fill[level]
    vals = [level, risk, mitigation, risk_impacts[ri]]
    for ci, v in enumerate(vals):
        c = row.cells[ci]
        cell_bg(c, bg if ci==0 else ('FFFFFF' if ri%2==0 else 'FAFAFA')); cell_border(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci==0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(v); r.font.size=Pt(8.5); r.bold=(ci==0)
        r.font.color.rgb = RGBColor(int(fc[0:2],16),int(fc[2:4],16),int(fc[4:6],16)) if ci==0 else RGBColor(0x33,0x33,0x33)

# Footer
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run(
    f'Total Effort: {grand_pw} Person-Weeks  ({grand_pd} Person-Days)   |   '
    f'Project Duration: 16 Weeks (~4 Months)   |   Team: 5 members   |   '
    'SHARP EV AI Chatbot — Confidential')
r.font.size=Pt(8); r.font.italic=True
r.font.color.rgb=RGBColor(0x90,0x90,0x90)

output = '/home/ramram/Desktop/Personal/EV-Chatbot/SHARP_EV_Effort_Estimation.docx'
doc.save(output)
print(f'Saved: {output}')
print(f'Grand Total: {grand_pw} PW  /  {grand_pd} person-days')
print(f'Breakdown — SA:{tot_sa} | AIL:{tot_ail} | AIE1:{tot_aie1} | AIE2:{tot_aie2} | QA:{tot_qa}')
