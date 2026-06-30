"""
Generate SHARP EV AI Chatbot — AWS Production Cost Breakdown DOCX
Region: ap-northeast-1 (Tokyo)
Output: archive/documents/SHARP_EV_AWS_Production_Cost_Breakdown.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page setup: Landscape A4 ──────────────────────────────────────────────────
sec = doc.sections[0]
sec.orientation   = WD_ORIENT.LANDSCAPE
sec.page_width    = Cm(29.7)
sec.page_height   = Cm(21.0)
sec.top_margin    = Cm(1.5)
sec.bottom_margin = Cm(1.5)
sec.left_margin   = Cm(1.5)
sec.right_margin  = Cm(1.5)


# ── Helpers ───────────────────────────────────────────────────────────────────
def cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def cell_border(cell, color='000000', sz='4'):
    tc      = cell._tc
    tcPr    = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def write_cell(cell, text, bold=False, fs=9.0,
               align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    """Write text into a cell; handles \\n as XML line breaks."""
    lines = str(text).split('\n')
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    for idx, line in enumerate(lines):
        if idx > 0:
            run = p.add_run()
            br  = OxmlElement('w:br')
            run._r.append(br)
        r = p.add_run(line)
        r.bold   = bold
        r.italic = italic
        r.font.size      = Pt(fs)
        r.font.color.rgb = RGBColor(0, 0, 0)


def make_table(headers, rows, col_widths, hdr_fs=9.5, row_fs=9.0,
               center_cols=None):
    """Build a standard bordered table with grey header + alternating rows."""
    center_cols = center_cols or []
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style     = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_col_widths(tbl, col_widths)

    # Header
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        cell_bg(c, 'D9D9D9')
        cell_border(c, '000000', '6')
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        write_cell(c, h, bold=True, fs=hdr_fs,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg  = 'FFFFFF' if ri % 2 == 0 else 'F2F2F2'
        for ci, val in enumerate(row_data):
            c = row.cells[ci]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_bg(c, bg)
            cell_border(c, '000000', '4')
            align = (WD_ALIGN_PARAGRAPH.CENTER if ci in center_cols
                     else WD_ALIGN_PARAGRAPH.LEFT)
            write_cell(c, val, fs=row_fs, align=align)

    return tbl


def section_heading(text):
    h = doc.add_heading(text, 2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


# =============================================================================
# TITLE
# =============================================================================
t = doc.add_heading('SHARP EV AI Chatbot', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

s = doc.add_heading('AWS Production Cost Breakdown — ap-northeast-1 (Tokyo)', 2)
s.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run(
    'Prepared for: SHARP Business & Technical Team   |   '
    'Date: 2026-06-25   |   Version: 1.0   |   Confidential')
mr.font.size   = Pt(9)
mr.font.italic = True
doc.add_paragraph()


# =============================================================================
# PART 1 — SERVICE INVENTORY
# =============================================================================
section_heading('Part 1 — Complete Production Service Inventory  (20 Services)')

inv_headers = ['#', 'Layer', 'AWS Service', 'Purpose in Production', 'Paid?']
inv_data = [
    ('1',  'AI / Search',       'Amazon Kendra Enterprise',         'Document indexing, RAG retrieval, Japanese MeCab tokenizer',                'Yes'),
    ('2',  'AI / LLM',          'Amazon Bedrock (Claude 3 Haiku)',   'LLM answer generation from Kendra retrieved document excerpts',             'Yes'),
    ('3',  'AI / Voice',        'Amazon Transcribe',                 'Voice-to-Text (STT) — Japanese ja-JP batch transcription',                  'Yes'),
    ('4',  'AI / Voice',        'Amazon Polly',                      'Text-to-Voice (TTS) — Mizuki (JA Female Standard) / Takumi (JA Male Neural)','Yes'),
    ('5',  'AI / Image',        'Amazon Rekognition',                'Image label detection + OCR text detection from vehicle photos',             'Yes'),
    ('6',  'Storage',           'Amazon S3',                         'EV manuals, Kendra metadata JSON, Transcribe audio temp files',             'Yes'),
    ('7',  'Compute',           'Amazon EC2 Auto Scaling Group',     'Host Streamlit chatbot app — auto-scales across 2 Availability Zones',      'Yes'),
    ('8',  'Compute',           'Amazon ECR',                        'Docker container registry — stores versioned application images',            'Yes'),
    ('9',  'Network',           'Application Load Balancer (ALB)',   'HTTPS routing, SSL termination, health checks to app instances',            'Yes'),
    ('10', 'Network',           'Amazon VPC',                        'Private network isolation for all production resources',                    'Free'),
    ('11', 'Network',           'NAT Gateway',                       'Outbound internet access from private subnets to AWS APIs',                 'Yes'),
    ('12', 'Network',           'Amazon Route 53',                   'Custom domain DNS management, health-based routing for HA',                 'Yes'),
    ('13', 'Security',          'AWS Certificate Manager (ACM)',      'SSL/TLS certificates for HTTPS — auto-renewed at no cost',                  'Free'),
    ('14', 'Security',          'AWS WAF',                           'Web application firewall — blocks SQL injection, XSS, bot traffic',         'Yes'),
    ('15', 'Security',          'AWS Secrets Manager',               'Secure vault for API keys, Bedrock credentials, DB passwords',              'Yes'),
    ('16', 'Auth',              'Amazon Cognito',                    'User authentication — COCORO MEMBERS ID integration, JWT tokens',           'Partial'),
    ('17', 'Database',          'Amazon DynamoDB',                   'Chat history, user sessions, query audit logs — on-demand capacity',        'Yes'),
    ('18', 'Monitoring',        'Amazon CloudWatch',                 'Logs, metrics, dashboards, alarms for all services',                        'Partial'),
    ('19', 'Compliance',        'AWS CloudTrail',                    'Full API audit trail — mandatory for J-SOX and ISO 27001 in Japan',         'Partial'),
    ('20', 'Identity',          'AWS IAM',                           'Roles and policies for all service-to-service permissions',                 'Free'),
]
make_table(inv_headers, inv_data, [0.55, 2.3, 4.0, 10.5, 1.3],
           center_cols=[0, 4])
doc.add_paragraph()


# =============================================================================
# PART 2 — USAGE ASSUMPTIONS
# =============================================================================
section_heading('Part 2 — Monthly Usage Assumptions per Production Tier')

assump_headers = ['Parameter', 'Small Production', 'Medium Production', 'Enterprise Production']
assump_data = [
    ('Users per day',                    '200',                              '1,000',                      '5,000'),
    ('Active days / month',              '25',                               '25',                         '25'),
    ('Queries per user session (avg)',   '5',                                '5',                          '5'),
    ('Total queries / month',            '25,000',                           '125,000',                    '625,000'),
    ('Voice queries  (10% of total)',    '2,500  →  1,875 min',              '12,500  →  9,375 min',       '62,500  →  46,875 min'),
    ('Image queries  (10% of total)',    '2,500  →  5,000 API calls',        '12,500  →  25,000 API calls','62,500  →  125,000 API calls'),
    ('Avg voice duration',               '45 seconds',                       '45 seconds',                 '45 seconds'),
    ('Avg Bedrock input tokens',         '3,000 tokens',                     '3,000 tokens',               '3,000 tokens'),
    ('Avg Bedrock output tokens',        '700 tokens',                       '700 tokens',                 '700 tokens'),
    ('TTS responses  (20% of queries)',  '5,000 responses  →  3M chars',     '25,000  →  15M chars',       '125,000  →  75M chars'),
    ('Polly voice split',                '60% Mizuki (Std) / 40% Takumi (Neural)', 'Same',                 'Same'),
    ('App server config',                '2× EC2 t3.large',                  '3× EC2 t3.xlarge',           '4× EC2 t3.2xlarge'),
    ('Monthly Active Users (MAU)',       '~4,000',                           '~20,000',                    '~100,000'),
    ('Kendra extra query capacity',      'None  (within 8K QCU/day base)',   'None  (within 8K QCU/day)', '17,000 extra QCU/day'),
]
make_table(assump_headers, assump_data, [5.5, 5.2, 5.2, 5.2],
           hdr_fs=9.5, row_fs=9.0, center_cols=[])
doc.add_paragraph()
doc.add_page_break()


# =============================================================================
# PART 3 — MONTHLY COST BREAKDOWN  (custom table with summary rows)
# =============================================================================
section_heading('Part 3 — Detailed Monthly Cost Breakdown  (Tokyo Pricing — ap-northeast-1)')

# Column headers
C_HEADERS = [
    '#',
    'AWS Service',
    'Service Description / Purpose',
    'Unit Price  (Tokyo)',
    'Small Prod\n200 users/day',
    'Medium Prod\n1,000 users/day',
    'Enterprise\n5,000 users/day',
]
C_WIDTHS = [0.55, 3.5, 6.5, 4.3, 2.7, 2.7, 2.7]   # total ≈ 22.95 cm

# Row data: (#, Service, Description, Unit Price, Small $, Medium $, Enterprise $)
cost_rows = [
    ('1',  'Amazon Kendra\nEnterprise Edition',
           'Document RAG search & retrieval\nJapanese MeCab morphological tokenizer\nExtra QCU billed above 8,000 QCU/day base',
           '$1,008.00 / mo  (fixed)\n+ $0.00025 per extra QCU',
           '$1,008.00', '$1,008.00', '$1,135.50'),

    ('2',  'Amazon Bedrock\n(Claude 3 Haiku)',
           'LLM answer generation from Kendra excerpts\nInput tokens: system prompt + context + query\nOutput tokens: generated answer',
           '$0.00025 per 1K input tokens\n$0.00125 per 1K output tokens',
           '$40.63', '$203.13', '$1,015.63'),

    ('3',  'Amazon Transcribe\n(ja-JP)',
           'Voice-to-Text (STT)\nJapanese ja-JP batch transcription\nUploads audio to S3 → transcription job',
           '$0.024 per minute',
           '$45.00', '$225.00', '$1,125.00'),

    ('4',  'Amazon Polly\n(Mizuki + Takumi)',
           'Text-to-Voice (TTS)\nMizuki = Standard JA Female voice\nTakumi = Neural JA Male voice  (4× cost of Standard)',
           'Standard (Mizuki): $0.000004/char\nNeural (Takumi):   $0.000016/char',
           '$26.40', '$132.00', '$660.00'),

    ('5',  'Amazon Rekognition',
           'Image analysis for Self-Diagnosis module\nDetectLabels: identifies vehicle components\nDetectText: reads DTC fault codes from photos',
           '$0.001 per image  (per API call)\n2 API calls per image uploaded',
           '$5.00', '$25.00', '$125.00'),

    ('6',  'Amazon S3\n(Standard)',
           'Primary document + metadata store\nEV manuals, Kendra .metadata.json files\nTranscribe audio temp files',
           '$0.025 / GB / month\n$0.0053 per 1K PUT requests',
           '$0.67', '$3.34', '$16.74'),

    ('7',  'Amazon EC2\n(Auto Scaling Group)',
           'Hosts the Streamlit chatbot application\nDeployed across 2 Availability Zones for HA\nAuto-scales based on CPU / request load',
           'Small:  t3.large  ×2  @ $0.0832/hr\nMedium: t3.xlarge ×3  @ $0.1664/hr\nEnterprise: t3.2xlarge ×4 @ $0.3328/hr',
           '$121.47', '$364.52', '$972.58'),

    ('8',  'Amazon ECR',
           'Docker container registry\nStores versioned application container images\nUsed by EC2 Auto Scaling to pull latest image',
           '$0.10 per GB / month',
           '$0.50', '$1.00', '$2.00'),

    ('9',  'Application Load\nBalancer (ALB)',
           'HTTPS traffic routing to app instances\nSSL/TLS termination with ACM certificate\nHealth checks + sticky sessions',
           '$0.0243 / hr\n+ $0.008 per LCU-hour',
           '$22.74', '$37.74', '$47.74'),

    ('10', 'Amazon VPC',
           'Private network isolation\nPublic + private subnets in 2 AZs\nSecurity groups for all services',
           'Free',
           '$0.00', '$0.00', '$0.00'),

    ('11', 'NAT Gateway',
           'Outbound internet from private subnets\nRequired for EC2 → Kendra / Bedrock API calls\n1 NAT per AZ for HA',
           '$0.045 / hr\n+ $0.045 per GB data processed',
           '$33.30', '$35.00', '$41.85'),

    ('12', 'Amazon Route 53',
           'Custom domain DNS management\nLatency-based routing for Japan users\nHealth checks tied to ALB targets',
           '$0.50 / hosted zone / month\n+ $0.40 per 1M DNS queries',
           '$0.90', '$2.50', '$5.00'),

    ('13', 'AWS Certificate\nManager (ACM)',
           'SSL/TLS certificate for HTTPS\nAuto-renewed — no manual renewal needed\nAttached directly to ALB',
           'Free  (public certificates)',
           '$0.00', '$0.00', '$0.00'),

    ('14', 'AWS WAF',
           'Web application firewall on ALB\nBlocks SQL injection, XSS, bot traffic\nOWASP Top 10 managed rule groups',
           '$5.00 / WebACL / month\n+ $1.00 per 1M requests',
           '$10.00', '$9.00', '$23.75'),

    ('15', 'AWS Secrets Manager',
           'Secure vault for sensitive credentials\nBedrock API key, Kendra role ARN\nDynamoDB connection, Cognito client secret',
           '$0.40 / secret / month\n+ $0.05 per 10K API calls',
           '$2.05', '$2.05', '$2.05'),

    ('16', 'Amazon Cognito',
           'User authentication and session management\nCOCORO MEMBERS ID integration via OAuth2\nJWT token issuance and validation',
           'Free  ≤ 50,000 MAU\n$0.0055 / MAU  above 50K',
           '$0.00', '$0.00', '$275.00'),

    ('17', 'Amazon DynamoDB',
           'Chat history and user session store\nOn-demand read/write capacity units\nStores conversation context per user',
           '$1.25 per 1M write capacity units\n$0.25 per 1M read capacity units\n$0.285 per GB storage / month',
           '$0.58', '$4.50', '$85.00'),

    ('18', 'Amazon CloudWatch',
           'Centralised logs, metrics and alarms\nKendra sync job logs, app error tracking\nDashboard for real-time monitoring',
           '$0.76 per GB log ingestion\n$0.033 per GB log storage / month\n$0.30 per custom metric / month',
           '$10.00', '$25.00', '$80.00'),

    ('19', 'AWS CloudTrail',
           'Full API audit trail for all 20 services\nMandatory for J-SOX compliance in Japan\nRequired for ISO 27001 certification',
           'Free  (management events, 1 trail)\n$2.00 per 100K data events',
           '$2.00', '$5.00', '$20.00'),

    ('20', 'AWS IAM',
           'Roles and inline policies\nKendra execution role, Bedrock access policy\nS3 read permissions, EC2 instance profile',
           'Free',
           '$0.00', '$0.00', '$0.00'),
]

TOTAL_ROWS = len(cost_rows)   # 20
tbl = doc.add_table(rows=1 + TOTAL_ROWS + 2, cols=7)
tbl.style     = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_col_widths(tbl, C_WIDTHS)

# Header row
hrow = tbl.rows[0]
for i, h in enumerate(C_HEADERS):
    c = hrow.cells[i]
    cell_bg(c, 'D9D9D9')
    cell_border(c, '000000', '6')
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    write_cell(c, h, bold=True, fs=9.0, align=WD_ALIGN_PARAGRAPH.CENTER)

# Data rows
for ri, rd in enumerate(cost_rows):
    row = tbl.rows[ri + 1]
    bg  = 'FFFFFF' if ri % 2 == 0 else 'F2F2F2'
    for ci, val in enumerate(rd):
        c = row.cells[ci]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_bg(c, bg)
        cell_border(c, '000000', '4')
        align = (WD_ALIGN_PARAGRAPH.CENTER if ci in [0, 4, 5, 6]
                 else WD_ALIGN_PARAGRAPH.LEFT)
        write_cell(c, val, bold=(ci == 0), fs=8.5, align=align)

# Monthly total row
def fill_summary_row(row, label, small, medium, enterprise, bg='D9D9D9', sz='6'):
    vals = ['', label, '', '', small, medium, enterprise]
    for ci, val in enumerate(vals):
        c = row.cells[ci]
        cell_bg(c, bg)
        cell_border(c, '000000', sz)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        align = (WD_ALIGN_PARAGRAPH.CENTER if ci in [0, 4, 5, 6]
                 else WD_ALIGN_PARAGRAPH.LEFT)
        write_cell(c, val, bold=True, fs=9.5, align=align)

fill_summary_row(tbl.rows[TOTAL_ROWS + 1],
                 'TOTAL  /  MONTH', '$1,329', '$2,083', '$5,633')
fill_summary_row(tbl.rows[TOTAL_ROWS + 2],
                 'TOTAL  /  YEAR  (× 12)', '$15,948', '$24,996', '$67,596',
                 bg='F2F2F2')

doc.add_paragraph()
doc.add_page_break()


# =============================================================================
# PART 4 — COST BY CATEGORY
# =============================================================================
section_heading('Part 4 — Monthly Cost by Service Category')

cat_headers = ['Category', 'Services Included',
               'Small Prod\n($)', 'Medium Prod\n($)', 'Enterprise\n($)',
               '% of Total\n(Small Prod)']
cat_data = [
    ('AI / ML Core',            'Kendra + Bedrock + Transcribe + Polly + Rekognition', '$1,125',  '$1,593',  '$4,061',  '85%'),
    ('Compute',                 'EC2 Auto Scaling + ECR',                               '$122',    '$366',    '$975',    '9%'),
    ('Network',                 'ALB + NAT Gateway + Route 53',                         '$57',     '$75',     '$95',     '4%'),
    ('Storage & Database',      'S3 + DynamoDB',                                        '$1',      '$8',      '$102',    '<1%'),
    ('Security & Auth',         'WAF + Secrets Manager + Cognito + ACM',                '$12',     '$11',     '$301',    '1%'),
    ('Monitoring & Compliance', 'CloudWatch + CloudTrail',                              '$12',     '$30',     '$100',    '1%'),
    ('Identity',                'IAM',                                                  '$0',      '$0',      '$0',      '0%'),
    ('TOTAL',                   '20 Services',                                          '$1,329',  '$2,083',  '$5,633',  '100%'),
]
make_table(cat_headers, cat_data, [3.5, 7.2, 2.4, 2.4, 2.4, 2.4],
           center_cols=[2, 3, 4, 5])
doc.add_paragraph()


# =============================================================================
# PART 5 — KEY COST DRIVERS
# =============================================================================
section_heading('Part 5 — Key Cost Drivers by Production Tier')

drv_headers = ['Rank', 'AWS Service', 'Small Prod', 'Medium Prod', 'Enterprise', 'Why It Matters']
drv_data = [
    ('#1', 'Amazon Kendra Enterprise',
     '$1,008  (76%)', '$1,008  (48%)', '$1,136  (20%)',
     'Fixed monthly cost regardless of query volume.\nBreak-even vs. per-query pricing at ~1,000 queries/month.'),
    ('#2', 'Amazon EC2  (App Servers)',
     '$121  (9%)', '$365  (17%)', '$973  (17%)',
     'Fixed compute cost tied to instance size.\nUse schedule-based Auto Scaling to reduce off-peak cost by 30–40%.'),
    ('#3', 'Amazon Bedrock  (Claude 3 Haiku)',
     '$41  (3%)', '$203  (10%)', '$1,016  (18%)',
     'Very low at small scale. Scales with query volume.\nProvision Throughput mode reduces cost by ~50% above 100K queries/month.'),
    ('#4', 'Amazon Transcribe  (ja-JP)',
     '$45  (3%)', '$225  (11%)', '$1,125  (20%)',
     'Becomes the largest variable cost at enterprise scale.\nConsider caching transcripts for repeated voice queries.'),
    ('#5', 'Amazon Polly  (Takumi Neural)',
     '$26  (2%)', '$132  (6%)', '$660  (12%)',
     'Neural (Takumi) costs 4× more than Standard (Mizuki).\nSwitch non-critical TTS to Mizuki Standard to cut Polly cost by 75%.'),
]
make_table(drv_headers, drv_data, [1.0, 3.5, 2.5, 2.5, 2.5, 8.0],
           center_cols=[0, 2, 3, 4])
doc.add_paragraph()


# =============================================================================
# PART 6 — PRODUCTION NOTES FOR JAPAN
# =============================================================================
section_heading('Part 6 — Important Notes for Japan Production Deployment')

notes_headers = ['Topic', 'Detail']
notes_data = [
    ('Data Residency',
     'All 20 services are configured to ap-northeast-1 (Tokyo). No data — including voice, images, or chat history — leaves Japan.'),
    ('J-SOX Compliance',
     'AWS CloudTrail is mandatory for all API audit logging in Japanese enterprise environments. Required for J-SOX and ISO 27001 certification.'),
    ('Cognito Free Tier',
     'First 50,000 Monthly Active Users (MAU) are completely free. Small and Medium Production tiers are fully within the Cognito free tier.'),
    ('Polly Cost Control',
     'Switching non-critical responses from Takumi (Neural, $0.000016/char) to Mizuki (Standard, $0.000004/char) reduces Polly cost by 75%.'),
    ('Transcribe at Scale',
     'Transcribe becomes the #1 variable cost driver at 5,000+ users/day. Caching and deduplication of repeated voice queries can reduce cost significantly.'),
    ('EC2 Auto Scaling',
     'Schedule-based scaling (scale down 9pm–8am JST and weekends) reduces EC2 cost by approximately 30–40% without impacting business hours.'),
    ('NAT Gateway Optimisation',
     'Configure VPC Endpoints for S3, DynamoDB, and Kendra. Bypasses NAT Gateway, reduces data charges, and improves latency.'),
    ('Kendra Enterprise QCU Limit',
     'Enterprise Edition includes 8,000 Query Capacity Units per day. At 5,000 users × 5 queries = 25,000 queries/day, extra QCU charges apply ($127.50/month).'),
    ('Bedrock Provisioned Throughput',
     'For Enterprise Production (625K+ queries/month), Provisioned Throughput for Claude 3 Haiku can reduce per-token cost by up to 50%.'),
    ('AWS EDP / JDM Discount',
     'AWS Enterprise Discount Program or Japan Direct Market pricing reduces total invoices by 15–30% for 1- or 3-year annual commitments.'),
]
make_table(notes_headers, notes_data, [4.5, 16.5])
doc.add_paragraph()


# ── Footer ────────────────────────────────────────────────────────────────────
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run(
    'SHARP EV AI Chatbot  |  AWS Production Cost Breakdown  |  ap-northeast-1 (Tokyo)  |  '
    'Prices based on AWS public list pricing as of June 2026. '
    'Actual costs may vary based on usage patterns and negotiated discounts.  |  Confidential')
fr.font.size   = Pt(7.5)
fr.font.italic = True


# ── Save ──────────────────────────────────────────────────────────────────────
OUTPUT = ('/home/ramram/Desktop/Personal/EV-Chatbot/archive/documents/'
          'SHARP_EV_AWS_Production_Cost_Breakdown.docx')
doc.save(OUTPUT)
print(f'Saved: {OUTPUT}')
