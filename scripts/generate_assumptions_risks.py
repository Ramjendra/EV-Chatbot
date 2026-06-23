from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
sec = doc.sections[0]
sec.top_margin    = Cm(1.8)
sec.bottom_margin = Cm(1.8)
sec.left_margin   = Cm(2.0)
sec.right_margin  = Cm(2.0)

# ── Helpers ───────────────────────────────────────────────────────────────────
def cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cell_border(cell, color='BFBFBF', sz='4'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

def add_cell_text(cell, text, bold=False, fs=9, color='1A1A1A',
                  align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.size = Pt(fs)
    r.font.color.rgb = RGBColor(
        int(color[0:2],16), int(color[2:4],16), int(color[4:6],16))

# ── Colours ───────────────────────────────────────────────────────────────────
HEADER_BG   = '1F497D'
TITLE_COL   = 'FFFFFF'
ROW_COLORS  = [
    ('1565C0', 'E3F2FD'),   # 1 blue
    ('C62828', 'FFEBEE'),   # 2 red
    ('6A1B9A', 'EDE7F6'),   # 3 purple
    ('E65100', 'FFF3E0'),   # 4 orange
    ('2E7D32', 'E8F5E9'),   # 5 green
]

# =============================================================================
# TITLE BLOCK
# =============================================================================
title = doc.add_heading('SHARP EV AI Chatbot', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_heading('Top 5 Assumptions, Risks & Dependencies', 2)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(
    'Prepared for: SHARP Business & Client Team   |   '
    'Date: 2026-06-23   |   Version: 1.0')
r.font.size = Pt(8.5)
r.font.italic = True
r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()

# ── Purpose note ──────────────────────────────────────────────────────────────
note_tbl = doc.add_table(rows=1, cols=1)
note_tbl.style = 'Table Grid'
nc = note_tbl.rows[0].cells[0]
cell_bg(nc, 'EBF3FB')
cell_border(nc, '2E74B5')
p = nc.paragraphs[0]
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(4)
r = p.add_run('Purpose: ')
r.bold = True; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r2 = p.add_run(
    'This document outlines the top 5 assumptions, risks, and dependencies '
    'that must be confirmed and managed to ensure successful delivery of the '
    'SHARP EV AI Chatbot project. Each item requires SHARP\'s acknowledgement '
    'and action before development begins.')
r2.font.size = Pt(9)

doc.add_paragraph()

# =============================================================================
# MAIN TABLE
# =============================================================================
data = [
    (
        '1',
        '📄  EV Manual\nQuality',
        'SHARP provides EV manuals in clean, readable digital format.',
        'Poor quality documents cause the AI to give wrong or incomplete answers to customers.',
        'Confirm document format before project starts. If scanned images, treat as a blocker before any AI development begins.',
    ),
    (
        '2',
        '🚗  Vehicle Sensor\nData Connection',
        'SHARP EV cars can send real-time data (battery, fault codes, warnings) to the cloud.',
        'Self-Diagnosis feature cannot be built. Feature will be removed from Version 1 and pushed to a later release.',
        'Confirm if vehicle data connection exists. If not available, Self-Diagnosis will be image-only in Version 1.',
    ),
    (
        '3',
        '🤖  AI Service\nApproval',
        'Amazon AI services (Claude AI, Titan Embeddings) are approved and active on AWS Japan before development starts.',
        'All 4 chatbot features are blocked at the same time. No backup option available in current design.',
        'Submit AI service approval request at project start. Confirm activation before development begins.',
    ),
    (
        '4',
        '📋  Open RFQ\nValues',
        'SHARP confirms all unanswered values from the RFQ — response accuracy, speed, number of users, usage limits, answer length.',
        'System is built for the wrong scale. Rebuilding in final weeks causes delays and extra cost.',
        'All open values must be confirmed in writing by SHARP before development starts.',
    ),
    (
        '5',
        '🔐  COCORO Login\nSystem',
        'SHARP provides COCORO MEMBERS ID login system technical details and test access before integration begins.',
        'Chatbot cannot go live without a working login. All user testing is blocked. SHARP security team will not approve go-live.',
        'Share COCORO technical specification and test credentials at project start. This is a hard requirement for launch.',
    ),
]

COL_W   = [0.8, 2.8, 4.2, 4.2, 4.2]
HEADERS = ['#', 'Topic', 'Assumption', 'Risk if Not Met', 'Mitigation']

tbl = doc.add_table(rows=1 + len(data), cols=5)
tbl.style       = 'Table Grid'
tbl.alignment   = WD_TABLE_ALIGNMENT.CENTER
set_col_widths(tbl, COL_W)

# ── Header row ────────────────────────────────────────────────────────────────
hrow = tbl.rows[0]
for i, h in enumerate(HEADERS):
    c = hrow.cells[i]
    cell_bg(c, HEADER_BG)
    cell_border(c, HEADER_BG, '6')
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_cell_text(c, h, bold=True, fs=10, color=TITLE_COL,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Data rows ─────────────────────────────────────────────────────────────────
for ri, (num, topic, assumption, risk, mitigation) in enumerate(data):
    dark, light = ROW_COLORS[ri]
    row = tbl.rows[ri + 1]
    row.height_rule = None

    vals = [num, topic, assumption, risk, mitigation]
    for ci, val in enumerate(vals):
        c = row.cells[ci]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Number cell — dark colour
        if ci == 0:
            cell_bg(c, dark)
            cell_border(c, dark, '6')
            add_cell_text(c, val, bold=True, fs=14, color='FFFFFF',
                          align=WD_ALIGN_PARAGRAPH.CENTER)

        # Topic cell — dark colour
        elif ci == 1:
            cell_bg(c, dark)
            cell_border(c, dark, '4')
            add_cell_text(c, val, bold=True, fs=9.5, color='FFFFFF',
                          align=WD_ALIGN_PARAGRAPH.CENTER)

        # Risk cell — light background, coloured text
        elif ci == 3:
            cell_bg(c, 'FFF8F8')
            cell_border(c, 'DDDDDD', '2')
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            r = p.add_run('⚠  ')
            r.font.size = Pt(9); r.bold = True
            r.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
            r2 = p.add_run(val)
            r2.font.size = Pt(9)
            r2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

        # Mitigation cell — light background, green text prefix
        elif ci == 4:
            cell_bg(c, 'F6FFF6')
            cell_border(c, 'DDDDDD', '2')
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            r = p.add_run('✔  ')
            r.font.size = Pt(9); r.bold = True
            r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            r2 = p.add_run(val)
            r2.font.size = Pt(9)
            r2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

        # Assumption cell — light bg
        else:
            cell_bg(c, light)
            cell_border(c, 'DDDDDD', '2')
            add_cell_text(c, val, fs=9)

doc.add_paragraph()

# =============================================================================
# ACTION REQUIRED FOOTER TABLE
# =============================================================================
h2 = doc.add_heading('Action Required from SHARP', 2)
h2.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

actions = [
    ('1', 'EV Manuals',           'Provide digital PDF files with readable text'),
    ('2', 'Vehicle Data',         'Confirm if vehicle sensor data connection exists and share details'),
    ('3', 'AI Service Approval',  'Submit and confirm AWS Japan AI service activation'),
    ('4', 'Open RFQ Values',      'Provide written confirmation of all unanswered items'),
    ('5', 'COCORO Login Details', 'Share COCORO MEMBERS ID specification and test credentials'),
]

act_tbl = doc.add_table(rows=1 + len(actions), cols=3)
act_tbl.style     = 'Table Grid'
act_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_col_widths(act_tbl, [0.8, 4.0, 11.4])

# Header
for i, h in enumerate(['#', 'Item', 'Action Required']):
    c = act_tbl.rows[0].cells[i]
    cell_bg(c, '2E4057')
    cell_border(c, '2E4057', '6')
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_cell_text(c, h, bold=True, fs=9.5, color='FFFFFF',
                  align=WD_ALIGN_PARAGRAPH.CENTER)

# Rows
for ri, (num, item, action) in enumerate(actions):
    dark, light = ROW_COLORS[ri]
    row = act_tbl.rows[ri + 1]
    for ci, val in enumerate([num, item, action]):
        c = row.cells[ci]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if ci == 0:
            cell_bg(c, dark); cell_border(c, dark, '4')
            add_cell_text(c, val, bold=True, fs=11, color='FFFFFF',
                          align=WD_ALIGN_PARAGRAPH.CENTER)
        elif ci == 1:
            cell_bg(c, dark); cell_border(c, dark, '4')
            add_cell_text(c, val, bold=True, fs=9, color='FFFFFF')
        else:
            cell_bg(c, 'FFFFFF' if ri % 2 == 0 else 'F8F9FA')
            cell_border(c, 'DDDDDD', '2')
            add_cell_text(c, val, fs=9)

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run(
    'All 5 items must be confirmed before development begins  '
    '|  SHARP EV AI Chatbot  |  Confidential')
r.font.size = Pt(8)
r.font.italic = True
r.font.color.rgb = RGBColor(0x90, 0x90, 0x90)

# ── Save ──────────────────────────────────────────────────────────────────────
output = '/home/ramram/Desktop/Personal/EV-Chatbot/archive/documents/SHARP_EV_Assumptions_Risks_Dependencies.docx'
doc.save(output)
print(f'Saved: {output}')
