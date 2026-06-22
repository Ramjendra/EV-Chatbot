from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color='BFBFBF'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

# ── Title block ──────────────────────────────────────────────────────────────
title = doc.add_heading('AI Chatbot for SHARP EV', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_heading('Requirements Clarification Questions', 2)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run('Document prepared by: Development Team   |   Date: 2026-06-16   |   Version: 1.0')
r.font.size = Pt(9)
r.font.italic = True
r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()

# ── Purpose note ─────────────────────────────────────────────────────────────
purpose_table = doc.add_table(rows=1, cols=1)
purpose_table.style = 'Table Grid'
c = purpose_table.rows[0].cells[0]
set_cell_bg(c, 'EBF3FB')
set_cell_border(c, '2E74B5')
p = c.paragraphs[0]
r = p.add_run('Purpose: ')
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r.font.size = Pt(9)
r2 = p.add_run(
    'The following questions are intended to be shared with SHARP to gather all necessary '
    'information before development begins. Answers will directly shape architecture decisions, '
    'technology choices, compliance requirements, and delivery scope.'
)
r2.font.size = Pt(9)

doc.add_paragraph()

# ── Categories and questions ──────────────────────────────────────────────────
categories = [
    {
        "number": "1",
        "title": "Self-Diagnosis",
        "color": "1F497D",
        "text_color": RGBColor(0x1F, 0x49, 0x7D),
        "questions": [
            "What vehicle data is uploaded to the cloud — fault codes (OBD-II/DTC), sensor telemetry, battery status, or all of the above?",
            "What is the data protocol/format from the vehicle (proprietary API, CAN bus, REST endpoint)?",
            "How real-time must the vehicle data be for diagnosis — live streaming vs. snapshot on request?",
            "For image input: what are users expected to photograph — warning lights, physical damage, error screens?",
            "Should diagnosis output include severity levels (critical / warning / info) and suggested actions (visit dealer, safe to drive, etc.)?",
            "Is the diagnosis intended for end users (simple language) or service technicians (technical codes)?",
        ]
    },
    {
        "number": "2",
        "title": "EV Usage Consultation",
        "color": "2E74B5",
        "text_color": RGBColor(0x2E, 0x74, 0xB5),
        "questions": [
            "Do you have Sharp EV manuals in a structured digital format (PDF, HTML, JSON)? If PDF, how many pages/models?",
            "Which EV models and variants must be supported in v1?",
            "How often are manuals updated, and who is responsible for keeping the knowledge base current?",
            "What should happen when the manual has no answer — fallback to general guidance, say \"contact support\", or escalate?",
        ]
    },
    {
        "number": "3",
        "title": "Chat (General AI)",
        "color": "20695E",
        "text_color": RGBColor(0x20, 0x69, 0x5E),
        "questions": [
            "Which LLM provider is preferred — OpenAI, Gemini, Copilot, or open to recommendation?",
            "Are there data privacy restrictions on sending user queries to external LLM APIs (GDPR, Japan PIPA)?",
            "Should conversation history be persisted across sessions? If so, for how long?",
            "Are there any topic guardrails — topics the chatbot must refuse to answer?",
        ]
    },
    {
        "number": "4",
        "title": "General Automotive Guidance",
        "color": "7B2C9E",
        "text_color": RGBColor(0x7B, 0x2C, 0x9E),
        "questions": [
            "What knowledge base backs this — public automotive standards, internal SHARP documentation, or a third-party data source?",
            "Is scope limited to EVs only, or all vehicle types?",
        ]
    },
    {
        "number": "5",
        "title": "Platform & Integration",
        "color": "C5500C",
        "text_color": RGBColor(0xC5, 0x50, 0x0C),
        "questions": [
            "What is the target platform — mobile app (iOS/Android), web browser, in-car display, or all?",
            "Is there an existing SHARP app/portal this chatbot must embed into, or is it a standalone product?",
            "Does the chatbot need to integrate with dealer booking / service center systems?",
            "Is user authentication required? If yes, what identity provider is used?",
        ]
    },
    {
        "number": "6",
        "title": "Voice I/O",
        "color": "833C00",
        "text_color": RGBColor(0x83, 0x3C, 0x00),
        "questions": [
            "For voice input/output, which TTS/STT engines are preferred or already licensed?",
            "What languages must be supported in v1? (Japanese only, or multilingual?)",
            "Should voice and text modes be available simultaneously or selectable?",
        ]
    },
    {
        "number": "7",
        "title": "Non-Functional",
        "color": "176D6D",
        "text_color": RGBColor(0x17, 0x6D, 0x6D),
        "questions": [
            "What are the expected concurrent users and SLA (uptime %, response time targets)?",
            "Are there accessibility requirements (WCAG, screen reader support)?",
            "What analytics or conversation logging is required — full transcripts, intent tracking, error rates?",
            "Is there a staging/test environment with real vehicle data, or will mock data be provided for development?",
        ]
    },
]

total_questions = 0

for cat in categories:
    # Category header table (full-width colored band)
    hdr_table = doc.add_table(rows=1, cols=1)
    hdr_table.style = 'Table Grid'
    hdr_cell = hdr_table.rows[0].cells[0]
    set_cell_bg(hdr_cell, cat["color"])
    p = hdr_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(f'  {cat["number"]}.  {cat["title"]}')
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(11)

    # Questions table
    q_table = doc.add_table(rows=len(cat["questions"]), cols=2)
    q_table.style = 'Table Grid'

    col_widths = [Cm(1.5), Cm(14.5)]

    for q_idx, question in enumerate(cat["questions"]):
        row = q_table.rows[q_idx]
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]

        bg = 'F7FBFF' if q_idx % 2 == 0 else 'FFFFFF'

        for cell in row.cells:
            set_cell_bg(cell, bg)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Q number cell
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'Q{total_questions + q_idx + 1}')
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = cat["text_color"]

        # Question cell
        p = row.cells[1].paragraphs[0]
        r = p.add_run(question)
        r.font.size = Pt(9.5)

        # Answer space row
    # Add a blank answer-space row after each category's questions
    answer_table = doc.add_table(rows=1, cols=2)
    answer_table.style = 'Table Grid'
    answer_table.rows[0].cells[0].width = col_widths[0]
    answer_table.rows[0].cells[1].width = col_widths[1]
    set_cell_bg(answer_table.rows[0].cells[0], 'FFF9F0')
    set_cell_bg(answer_table.rows[0].cells[1], 'FFF9F0')
    set_cell_border(answer_table.rows[0].cells[0], 'E0C080')
    set_cell_border(answer_table.rows[0].cells[1], 'E0C080')
    p = answer_table.rows[0].cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Notes')
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xA0, 0x80, 0x40)
    answer_cell = answer_table.rows[0].cells[1].paragraphs[0]
    r2 = answer_cell.add_run(' ' * 5)
    r2.font.size = Pt(20)  # tall blank row for handwritten/typed notes

    total_questions += len(cat["questions"])
    doc.add_paragraph()

# ── Footer summary ────────────────────────────────────────────────────────────
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_p.add_run(
    f'Total: {total_questions} questions across {len(categories)} categories   |   '
    'SHARP EV AI Chatbot — Confidential'
)
r.font.size = Pt(8)
r.font.italic = True
r.font.color.rgb = RGBColor(0x90, 0x90, 0x90)

output_path = '/home/ramram/Desktop/Personal/EV-Chatbot/SHARP_EV_Requirements_Questions.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
