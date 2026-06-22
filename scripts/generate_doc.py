from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

# Title
title = doc.add_heading('SHARP EV AI Chatbot', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.runs[0]
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_heading('Client Clarification Questions', 2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph('')

questions = [
    ("Self-Diagnosis", "What vehicle data is uploaded to the cloud — fault codes (OBD-II/DTC), sensor telemetry, battery status, or all?", "Determines data pipeline and backend architecture"),
    ("Self-Diagnosis", "What is the data protocol/format from the vehicle (proprietary API, CAN bus, REST endpoint)?", "Defines integration layer design"),
    ("Self-Diagnosis", "How real-time must vehicle data be — live streaming vs. snapshot on request?", "Impacts infrastructure cost and complexity"),
    ("Self-Diagnosis", "For image input: what are users expected to photograph — warning lights, physical damage, error screens?", "Scopes the computer vision / image analysis feature"),
    ("Self-Diagnosis", "Should diagnosis output include severity levels (critical / warning / info) and suggested actions?", "Defines response format and UX flow"),
    ("Self-Diagnosis", "Is diagnosis intended for end users (simple language) or service technicians (technical codes)?", "Determines tone, depth, and content of responses"),
    ("Self-Diagnosis", "After diagnosis, can the chatbot trigger an action — book a service appointment or contact a dealer?", "Scopes integration with dealer/service systems"),
    ("Self-Diagnosis", "Should diagnosis history be stored per vehicle/user over time?", "Impacts database design and data retention policy"),
    ("Self-Diagnosis", "What fault types are in scope for v1 — battery, motor, charging, ADAS, or all systems?", "Defines KB and rules engine scope"),
    ("Self-Diagnosis", "Is there a severity threshold above which the bot must always recommend visiting a dealer?", "Defines escalation logic and safety guardrails"),
    ("EV Usage Consultation", "Do you have Sharp EV manuals in structured digital format (PDF, HTML, JSON)?", "Determines KB ingestion effort and RAG architecture"),
    ("EV Usage Consultation", "Which EV models and variants must be supported in v1?", "Scopes the knowledge base size"),
    ("EV Usage Consultation", "How often are manuals updated, and who is responsible for keeping the KB current?", "Defines content ops process and versioning strategy"),
    ("EV Usage Consultation", "What should happen when the manual has no answer — fallback, say 'contact support', or escalate?", "Defines fallback and escalation flow"),
    ("EV Usage Consultation", "Should answers reference specific page numbers or sections from the manual?", "Impacts citation/traceability feature requirement"),
    ("EV Usage Consultation", "How should the chatbot handle queries about features not yet in the current manual?", "Defines out-of-scope handling behavior"),
    ("General Chat", "Which LLM provider is preferred — OpenAI, Gemini, Copilot, or open to recommendation?", "Core technology decision affecting cost and capability"),
    ("General Chat", "Are there data privacy restrictions on sending user queries to external LLM APIs (GDPR, Japan PIPA)?", "May block use of external LLMs entirely"),
    ("General Chat", "Should conversation history be persisted across sessions? If yes, for how long?", "Impacts storage design and privacy compliance"),
    ("General Chat", "Are there topic guardrails — topics the chatbot must refuse to answer?", "Defines content moderation and safety layer"),
    ("General Chat", "Should general chat be strictly scoped to EV/automotive topics, or fully open-ended?", "Determines LLM system prompt and filtering logic"),
    ("General Chat", "Should the chatbot disclose to users when it is using a third-party AI service?", "Transparency/compliance requirement"),
    ("General Automotive Guidance", "What knowledge base backs this — public sources, internal SHARP docs, or third-party data?", "Defines data sourcing and licensing requirements"),
    ("General Automotive Guidance", "Is scope limited to EVs only, or all vehicle types?", "Scopes KB breadth"),
    ("Voice I/O", "Which TTS/STT engines are preferred or already licensed?", "Avoids vendor lock-in conflicts"),
    ("Voice I/O", "What languages must be supported in v1?", "Impacts TTS/STT model selection"),
    ("Voice I/O", "Should voice and text modes be available simultaneously or user-selectable?", "Defines UI/UX and accessibility design"),
    ("User & Personas", "Who are the primary users — EV owners, fleet managers, or dealership technicians?", "Shapes tone, depth, and feature priorities"),
    ("User & Personas", "Should the chatbot support multiple user roles with different access levels?", "Drives auth and role-based response logic"),
    ("Conversation Design", "Should the chatbot have a defined persona, name, or avatar aligned with the SHARP brand?", "Brand and UX requirement"),
    ("Conversation Design", "Should it support multi-turn conversations (remembering context within a session)?", "Core architectural decision"),
    ("Conversation Design", "How should the chatbot handle a query that spans multiple modules — e.g., usage question leads into diagnosis?", "Defines module-switching and context handoff logic"),
    ("Conversation Design", "Should it proactively suggest topics, or only respond to user-initiated queries?", "UX and engagement design decision"),
    ("Platform & Integration", "What is the target platform — mobile app, web browser, in-car display, or all?", "Core delivery scope decision"),
    ("Platform & Integration", "Is there an existing SHARP app/portal to embed into, or is this a greenfield product?", "Determines integration vs. build-from-scratch effort"),
    ("Platform & Integration", "What is the preferred cloud provider — AWS, Azure, Google Cloud, or on-premise?", "Infrastructure and compliance constraint"),
    ("Platform & Integration", "What is the expected number of active EV units at launch?", "Capacity planning and cost estimation"),
    ("Platform & Integration", "Should the system support offline or low-connectivity scenarios?", "Edge/offline architecture requirement"),
    ("Data & Privacy", "What user data can be stored — conversation logs, vehicle data, personal identifiers?", "Data architecture and retention policy"),
    ("Data & Privacy", "Must data reside in Japan only, or is global cloud storage acceptable?", "Data residency and compliance constraint"),
    ("Data & Privacy", "Are voice recordings allowed to be stored, or must they be discarded after transcription?", "Privacy and compliance requirement"),
    ("Data & Privacy", "Which compliance standards apply — Japan PIPA, GDPR, ISO 27001, UN R155?", "Non-negotiable legal requirements"),
    ("Handoff & Escalation", "Is there a human agent fallback — live chat or call center — when the bot fails?", "Defines escalation path and integration scope"),
    ("Handoff & Escalation", "Should full conversation history be passed to the human agent on handoff?", "CRM/helpdesk integration requirement"),
    ("Handoff & Escalation", "Should the bot detect user frustration and proactively offer escalation?", "Sentiment analysis feature scope"),
    ("Non-Functional", "What are the expected concurrent users and SLA targets (uptime %, response time)?", "Infrastructure sizing and SLA design"),
    ("Non-Functional", "What analytics or conversation logging is required — transcripts, intent tracking, error rates?", "Observability and reporting requirement"),
    ("Operations & Success", "What is the target go-live date, and is there a phased rollout (pilot → full launch)?", "Project planning and scope prioritization"),
    ("Operations & Success", "Who owns ongoing content updates post-launch?", "Defines operational handoff and support model"),
    ("Operations & Success", "What KPIs define success — resolution rate, CSAT, call center deflection, response latency?", "Defines acceptance criteria and measurement plan"),
    ("Future Scope", "Is video input (beyond still images) planned for a future phase?", "Prevents architecture decisions that block later expansion"),
    ("Future Scope", "Is in-car head unit integration planned, or mobile/web only for now?", "Scopes v1 delivery boundary clearly"),
    ("Future Scope", "Are predictive maintenance push notifications planned?", "Defines event-driven vs. request-response architecture"),
]

# Header colors per category
category_colors = {
    "Self-Diagnosis":             RGBColor(0x1F, 0x49, 0x7D),
    "EV Usage Consultation":      RGBColor(0x2E, 0x74, 0xB5),
    "General Chat":               RGBColor(0x20, 0x6D, 0x5E),
    "General Automotive Guidance":RGBColor(0x7B, 0x2C, 0x9E),
    "Voice I/O":                  RGBColor(0xC5, 0x50, 0x0C),
    "User & Personas":            RGBColor(0x83, 0x3C, 0x00),
    "Conversation Design":        RGBColor(0x37, 0x5A, 0x7F),
    "Platform & Integration":     RGBColor(0x1F, 0x49, 0x7D),
    "Data & Privacy":             RGBColor(0xC0, 0x39, 0x2B),
    "Handoff & Escalation":       RGBColor(0x8E, 0x44, 0xAD),
    "Non-Functional":             RGBColor(0x17, 0x6D, 0x6D),
    "Operations & Success":       RGBColor(0x27, 0xAE, 0x60),
    "Future Scope":               RGBColor(0x95, 0xA5, 0xA6),
}

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'BFBFBF')
        tcBorders.append(border)
    tcPr.append(tcBorders)

# Create table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.autofit = False

# Set column widths
col_widths = [Cm(1.2), Cm(4.5), Cm(8.5), Cm(4.8)]
for i, width in enumerate(col_widths):
    for cell in table.columns[i].cells:
        cell.width = width

# Header row
headers = ['#', 'Category', 'Question', 'Why It Matters']
hdr_row = table.rows[0]
for i, (cell, header) in enumerate(zip(hdr_row.cells, headers)):
    set_cell_bg(cell, '1F497D')
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(header)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)

# Data rows
current_category = None
row_num = 1

for idx, (category, question, why) in enumerate(questions):
    row = table.add_row()
    row.cells[0].width = col_widths[0]
    row.cells[1].width = col_widths[1]
    row.cells[2].width = col_widths[2]
    row.cells[3].width = col_widths[3]

    is_new_category = (category != current_category)
    bg = 'EBF3FB' if idx % 2 == 0 else 'FFFFFF'

    for cell in row.cells:
        set_cell_bg(cell, bg)
        set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Number
    p = row.cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(row_num))
    r.font.size = Pt(9)
    r.font.bold = True
    cat_color = category_colors.get(category, RGBColor(0x1F, 0x49, 0x7D))
    r.font.color.rgb = cat_color

    # Category
    p = row.cells[1].paragraphs[0]
    r = p.add_run(category)
    r.font.size = Pt(9)
    r.font.bold = is_new_category
    r.font.color.rgb = cat_color

    # Question
    p = row.cells[2].paragraphs[0]
    r = p.add_run(question)
    r.font.size = Pt(9)

    # Why
    p = row.cells[3].paragraphs[0]
    r = p.add_run(why)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    current_category = category
    row_num += 1

# Footer note
doc.add_paragraph('')
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.add_run(f'Total: {len(questions)} questions across {len(set(q[0] for q in questions))} categories  |  SHARP EV AI Chatbot — Confidential')
r.font.size = Pt(8)
r.font.italic = True
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

output_path = '/home/ramram/Desktop/Personal/EV-Chatbot/SHARP_EV_Chatbot_Questions.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
