from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

# ── Helpers ───────────────────────────────────────────────────────────────────
def cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cell_border(cell, color='BFBFBF'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_section_header(doc, number, title, bg_hex, text_rgb):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.rows[0].cells[0]
    cell_bg(c, bg_hex)
    p = c.paragraphs[0]
    r = p.add_run(f'  {number}.  {title}')
    r.bold = True
    r.font.color.rgb = RGBColor(*text_rgb)
    r.font.size = Pt(11)

def add_source_note(doc, source_text):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.rows[0].cells[0]
    cell_bg(c, 'FFF9EC')
    cell_border(c, 'E0A800')
    p = c.paragraphs[0]
    r1 = p.add_run('Slide reference: ')
    r1.bold = True; r1.font.size = Pt(8)
    r1.font.color.rgb = RGBColor(0x85, 0x60, 0x04)
    r2 = p.add_run(source_text)
    r2.font.size = Pt(8); r2.font.italic = True
    r2.font.color.rgb = RGBColor(0x85, 0x60, 0x04)

def add_col_headers(doc):
    col_widths = [Cm(1.2), Cm(9.2), Cm(5.8)]
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    for i, (w, h) in enumerate(zip(col_widths, ['#', 'Question', "Architect's Rationale"])):
        c = t.rows[0].cells[i]
        c.width = w
        cell_bg(c, '2E4057')
        cell_border(c, '2E4057')
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def add_questions(doc, questions, text_rgb, q_start):
    col_widths = [Cm(1.2), Cm(9.2), Cm(5.8)]
    t = doc.add_table(rows=len(questions), cols=3)
    t.style = 'Table Grid'
    for idx, (question, rationale) in enumerate(questions):
        row = t.rows[idx]
        for i, w in enumerate(col_widths):
            row.cells[i].width = w
        bg = 'F5F9FF' if idx % 2 == 0 else 'FFFFFF'
        for cell in row.cells:
            cell_bg(cell, bg)
            cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'Q{q_start + idx}')
        r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(*text_rgb)
        p = row.cells[1].paragraphs[0]
        p.add_run(question).font.size = Pt(9.5)
        p = row.cells[2].paragraphs[0]
        rr = p.add_run(rationale)
        rr.font.size = Pt(8.5); rr.font.italic = True
        rr.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
    return q_start + len(questions)

# ── TITLE ─────────────────────────────────────────────────────────────────────
h = doc.add_heading('SHARP EV AI Chatbot', 0)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

s = doc.add_heading('Solution Architect — Requirements Gathering Questionnaire', 2)
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
s.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

m = doc.add_paragraph()
m.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = m.add_run('Role: Solution Architect / Requirements Analyst   |   Date: 2026-06-16   |   Version: 2.0   |   Confidential')
mr.font.size = Pt(9); mr.font.italic = True
mr.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()

# Purpose box
pt = doc.add_table(rows=1, cols=1)
pt.style = 'Table Grid'
pc = pt.rows[0].cells[0]
cell_bg(pc, 'EBF3FB'); cell_border(pc, '2E74B5')
p = pc.paragraphs[0]
r = p.add_run('Objective: ')
r.bold = True; r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
r2 = p.add_run(
    'Consolidated questions derived from SHARP EV RFQ slide deck (pages 13–17). '
    'Similar topics have been merged for brevity. Goal: resolve all TBD values, confirm integration '
    'touchpoints, clarify performance thresholds, and define compliance boundaries before '
    'architecture scoping begins.'
)
r2.font.size = Pt(9)

doc.add_paragraph()
q = 1

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 1 — SCOPE & DELIVERY  (Slide 13)
# ─────────────────────────────────────────────────────────────────────────────
add_section_header(doc, 1, 'Scope & Delivery', '1F497D', (255, 255, 255))
add_source_note(doc, 'Slide 13 — Functional Overview')
add_col_headers(doc)
q = add_questions(doc, [
    (
        "Should all 4 modules (Self-diagnosis, EV Usage Consultation, Chat, General Automotive Guidance) "
        "launch together in v1, or is a phased rollout acceptable? Should they appear as one unified chatbot or separate selectable modes?",
        "Drives MVP scope, release plan, and UI navigation architecture"
    ),
    (
        "For Self-diagnosis: does the vehicle-to-cloud data pipeline (fault codes, telemetry) already exist as a live API, "
        "or does it need to be built? What data format/protocol is used?",
        "If the pipeline doesn't exist, it is the critical path blocker for this entire module"
    ),
    (
        "Which EV models and variants must be supported in v1, and is the SHARP EV manual available in "
        "machine-readable format (not scanned PDF)?",
        "Model count determines KB complexity; scanned PDFs add significant OCR pre-processing effort"
    ),
    (
        "For General Chat and General Automotive Guidance — should both be restricted to EV/automotive topics only, "
        "and is there an existing licensed automotive knowledge base, or must content be sourced from scratch?",
        "Topic scope defines LLM guardrails; KB availability determines build vs. buy decision"
    ),
    (
        "When a user query spans multiple modules (e.g., a usage question leads into fault diagnosis), "
        "how should the chatbot handle the transition — seamlessly or with an explicit mode switch?",
        "Defines context-sharing and cross-module handoff logic"
    ),
], (0x1F, 0x49, 0x7D), q)

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 2 — EXISTING ASSETS & REFERENCE PRODUCTS  (Slide 14)
# ─────────────────────────────────────────────────────────────────────────────
add_section_header(doc, 2, 'Existing Assets & Reference Products', '20695E', (255, 255, 255))
add_source_note(doc, 'Slide 14 — Reference examples: SHARP CookTalk, COCORO HOME AI, Google Gemini, OpenAI ChatGPT')
add_col_headers(doc)
q = add_questions(doc, [
    (
        "Can the backend AI infrastructure or SDK from COCORO HOME AI / CookTalk be reused for this project? "
        "What AI stack does it use (LLM provider, vector DB, orchestration layer)?",
        "Reusing existing infrastructure significantly cuts build time and cost"
    ),
    (
        "What worked well and what failed in CookTalk and COCORO HOME AI that should directly influence "
        "architecture or UX decisions on this project?",
        "Prevents repeating past mistakes; surfaces institutional knowledge upfront"
    ),
    (
        "Should the EV chatbot follow COCORO HOME AI's design language and avatar persona, "
        "or is a distinct SHARP EV brand identity and voice character required?",
        "Design system reuse vs. new design effort; affects TTS voice selection"
    ),
    (
        "Does SHARP have existing enterprise API contracts with OpenAI or Google? "
        "Are Gemini and ChatGPT shown as a quality benchmark to match, or as mandatory vendor choices?",
        "Existing contracts unlock pricing and DPA; vendor mandate vs. benchmark changes LLM selection freedom"
    ),
], (0x20, 0x69, 0x5E), q)

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3 — AUTHENTICATION, VOICE & UX  (Slide 15)
# ─────────────────────────────────────────────────────────────────────────────
add_section_header(doc, 3, 'Authentication, Voice & UX', 'C5500C', (255, 255, 255))
add_source_note(doc, 'Slide 15 — UX Requirements: COCORO MEMBERS ID, Voice Interaction, Processing indicator, TTS assets')
add_col_headers(doc)
q = add_questions(doc, [
    (
        "Please share the COCORO MEMBERS ID OAuth/SSO API specification. "
        "Is login mandatory for all features, or is guest/anonymous access permitted for limited use? "
        "What are the session token lifespan and refresh rules?",
        "Auth spec is required before any backend work begins; guest access affects onboarding design"
    ),
    (
        "Voice interaction should remain active for 'a certain period' (TBD) — what is the specific "
        "auto-deactivation timeout after silence? Should manual disable be a UI toggle, a voice command, or both?",
        "TBD value must be confirmed to implement the voice session timer correctly"
    ),
    (
        "For TTS voice response: what is the approved monthly budget, preferred voice character "
        "(gender, tone, Japanese/multilingual), and should it match the COCORO HOME AI voice? "
        "Are existing SHARP brand audio/animation assets available for the processing indicator?",
        "Budget determines TTS provider; brand assets avoid duplicating design work"
    ),
    (
        "When responding to a voice query, should the text transcript always be shown alongside the audio, "
        "or is audio-only acceptable?",
        "Accessibility and comprehension requirement — also affects UI layout"
    ),
], (0xC5, 0x50, 0x0C), q)

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 4 — SERVER FUNCTIONAL: TBD VALUES & SECURITY  (Slide 16)
# ─────────────────────────────────────────────────────────────────────────────
add_section_header(doc, 4, 'Server Functional — TBD Values & Security', '7B2C9E', (255, 255, 255))
add_source_note(doc, 'Slide 16 — Server Functional Requirement: Error handling, Data management, Security, Answer quality, Safeguards, Response design, Fault diagnosis, Speech recognition')
add_col_headers(doc)
q = add_questions(doc, [
    (
        "Error handling shows 'TBD' — please confirm: (a) max timeout for LLM response, "
        "(b) max timeout for vehicle data API, (c) number of retries and backoff strategy (fixed/exponential)?",
        "Two separate timeouts needed — LLM (typically 5–30s) vs. vehicle API (1–5s); retry count affects cost"
    ),
    (
        "Conversation history must be stored (date/time, content, success/failure, error reason) — "
        "what is the retention period, and who has access: user only, SHARP internal teams, or support staff?",
        "Retention period drives storage cost and Japan PIPA compliance; access control defines security model"
    ),
    (
        "Answer accuracy target is 'over XX%' and speech recognition is 'over XX%' — "
        "what are the specific percentage targets? How will accuracy be measured (human eval, automated benchmarks, user ratings)?",
        "These are key acceptance criteria — both are TBD values that must be confirmed to set QA benchmarks"
    ),
    (
        "For security: (a) should ALL non-Japan IPs be blocked, or only suspicious ones? "
        "(b) what is the one-time token validity window for replay attack prevention? "
        "(c) what request rate (per minute/hour) triggers an abnormal usage block?",
        "Blanket geo-block vs. anomaly detection are very different implementations; thresholds needed for rate-limiter"
    ),
    (
        "Response length is 'within XXX characters' (TBD) — what is the specific limit? "
        "For image responses with URLs, which domain should links point to (SHARP site, manual portal, dealer pages)?",
        "Character limit affects LLM prompt design and truncation logic; URL domain defines content sourcing"
    ),
    (
        "Fault diagnosis should 'minimize questions to users' — what is the maximum number of "
        "clarifying questions allowed before delivering a diagnosis? "
        "Who authors the fallback response variants, and how many variants per scenario?",
        "A specific cap (e.g., max 2 questions) is needed to implement the diagnostic flow; content ownership must be agreed"
    ),
], (0x7B, 0x2C, 0x9E), q)

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 5 — NON-FUNCTIONAL: PERFORMANCE, COST & COMPLIANCE  (Slide 17)
# ─────────────────────────────────────────────────────────────────────────────
add_section_header(doc, 5, 'Non-Functional — Performance, Cost & Compliance', '176D6D', (255, 255, 255))
add_source_note(doc, 'Slide 17 — Non-Functional Requirement: Performance, Concurrent Users, LLM, Ethics, Usage Limits, METI Guidelines')
add_col_headers(doc)
q = add_questions(doc, [
    (
        "Performance targets show 'XX seconds' — please confirm specific values for: "
        "(a) App → Server latency, (b) App → Full answer end-to-end. "
        "Is streaming output (text appears word-by-word) acceptable to improve perceived speed?",
        "Both are TBD values critical for infra sizing; streaming can reduce perceived latency for slow LLM responses"
    ),
    (
        "Maximum concurrent users is 'XXXX' — what is the specific number, and is it Japan-only or global? "
        "What is the approved monthly budget for infrastructure and LLM API costs?",
        "Concurrent users = #1 input for capacity planning; budget constrains architecture choices"
    ),
    (
        "Who from SHARP approves the final LLM model selection? "
        "Are there SHARP-internal AI ethics guidelines stricter than METI that must also be followed?",
        "LLM sign-off process prevents vendor disputes; internal guidelines may tighten guardrail design scope"
    ),
    (
        "Usage limit is 'XX times/month' (TBD) — what is the specific cap per user? "
        "Should users receive a warning before cutoff (e.g., at 80%)? "
        "Are safety-critical modules like Self-diagnosis exempt from this limit?",
        "All three are TBD values; throttling a safety/diagnosis feature raises liability concerns"
    ),
    (
        "Which version of METI AI Business Guidelines applies, and are there specific articles requiring "
        "documented compliance evidence? Is a formal legal sign-off required before go-live?",
        "Specific METI version and articles determine compliance sprint scope; legal gate affects project timeline"
    ),
    (
        "What is the target go-live date, and is there a pilot/UAT phase before full launch?",
        "Drives overall project timeline, release strategy, and sprint planning"
    ),
], (0x17, 0x6D, 0x6D), q)

doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 6 — OPEN ITEMS & MISSING INFORMATION
# ─────────────────────────────────────────────────────────────────────────────
add_section_header(doc, 6, 'Open Items & Missing Information', 'C0392B', (255, 255, 255))
add_source_note(doc, 'Gaps identified across slides 13–17 of the 18-page RFQ deck')
add_col_headers(doc)
q = add_questions(doc, [
    (
        "The RFQ deck is 18 pages — we have reviewed pages 13–17. "
        "Can you share the complete deck (pages 1–12 and page 18) for full context?",
        "Earlier slides likely contain project background, budget, timeline, and architecture constraints"
    ),
    (
        "Is there a deployment architecture diagram (cloud provider, region, network topology)? "
        "Are any vendor API contracts (cloud, LLM, vehicle telemetry) already signed?",
        "Signed contracts constrain or enable specific technical choices; region drives data residency compliance"
    ),
    (
        "Are there UI/UX mockups or wireframes for the chatbot beyond the reference screenshots shown?",
        "Without mockups, UI development requires a separate design discovery phase"
    ),
    (
        "Who is the SHARP-side decision-making authority for scope changes and approvals? "
        "What are the UAT acceptance criteria and pass/fail definitions for each module?",
        "Single point of sign-off prevents scope creep; UAT criteria define the delivery finish line"
    ),
], (0xC0, 0x39, 0x2B), q)

doc.add_paragraph()

# ── Footer ────────────────────────────────────────────────────────────────────
total_q = q - 1
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run(
    f'Total: {total_q} questions across 6 sections   |   '
    'Based on SHARP EV RFQ Slides 13–17   |   Confidential — Solution Architect Working Document   |   v2.0'
)
r.font.size = Pt(8); r.font.italic = True
r.font.color.rgb = RGBColor(0x90, 0x90, 0x90)

output = '/home/ramram/Desktop/Personal/EV-Chatbot/SHARP_EV_SA_Questionnaire.docx'
doc.save(output)
print(f'Saved: {output}  ({total_q} questions)')
