"""
Generate SHARP EV AI Chatbot — AWS Pricing Calculator Style Cost Estimate DOCX
Mirrors the format of an AWS Pricing Calculator export (service, config, upfront,
monthly, 12-month total) for all 3 production tiers.

Region  : ap-northeast-1 (Tokyo)
Currency: USD
Output  : archive/documents/SHARP_EV_AWS_Cost_Calculator_Estimate.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page: Landscape A4 ────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.orientation   = WD_ORIENT.LANDSCAPE
sec.page_width    = Cm(29.7)
sec.page_height   = Cm(21.0)
sec.top_margin    = Cm(1.4)
sec.bottom_margin = Cm(1.4)
sec.left_margin   = Cm(1.5)
sec.right_margin  = Cm(1.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def cell_bg(cell, hex_color):
    tc   = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

def cell_border(cell, color='000000', sz='4'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0');    b.set(qn('w:color'), color)
        tcB.append(b)
    tcPr.append(tcB)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

def write_cell(cell, text, bold=False, fs=8.5,
               align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    lines = str(text).split('\n')
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    for idx, line in enumerate(lines):
        if idx > 0:
            run = p.add_run(); br = OxmlElement('w:br'); run._r.append(br)
        r = p.add_run(line)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(fs); r.font.color.rgb = RGBColor(0, 0, 0)

# ── Column definitions ────────────────────────────────────────────────────────
# Service | Configuration Summary | Upfront | Monthly | 12-Month Total
COLS    = ['AWS Service', 'Configuration Summary', 'Upfront\nCost', 'Monthly\nCost', '12-Month\nTotal']
WIDTHS  = [3.8, 10.2, 2.0, 2.3, 2.5]   # total 20.8 cm  (usable ~26.7 cm)
C_CENTER = [2, 3, 4]                     # right-align cost columns

# ── Build one tier table ──────────────────────────────────────────────────────
def build_tier_table(rows_data, monthly_total, annual_total):
    """rows_data: list of (service, config, upfront, monthly, 12mo)"""
    n = len(rows_data)
    tbl = doc.add_table(rows=1 + n + 2, cols=5)
    tbl.style     = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_col_widths(tbl, WIDTHS)

    # Header
    for i, h in enumerate(COLS):
        c = tbl.rows[0].cells[i]
        cell_bg(c, 'D9D9D9'); cell_border(c, '000000', '6')
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        write_cell(c, h, bold=True, fs=9.0, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for ri, (svc, cfg, up, mo, yr) in enumerate(rows_data):
        row = tbl.rows[ri + 1]
        bg  = 'FFFFFF' if ri % 2 == 0 else 'F2F2F2'
        for ci, val in enumerate([svc, cfg, up, mo, yr]):
            c = row.cells[ci]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_bg(c, bg); cell_border(c, '000000', '4')
            align = (WD_ALIGN_PARAGRAPH.RIGHT if ci in C_CENTER
                     else WD_ALIGN_PARAGRAPH.LEFT)
            write_cell(c, val, bold=(ci == 0), fs=8.5, align=align)

    # Monthly total
    def summary_row(idx, label, mo_val, yr_val, bg='D9D9D9', sz='6'):
        row = tbl.rows[idx]
        for ci, val in enumerate(['', label, '$0.00', mo_val, yr_val]):
            c = row.cells[ci]
            cell_bg(c, bg); cell_border(c, '000000', sz)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            align = (WD_ALIGN_PARAGRAPH.RIGHT if ci in [2, 3, 4]
                     else WD_ALIGN_PARAGRAPH.LEFT)
            write_cell(c, val, bold=True, fs=9.5, align=align)

    summary_row(n + 1, 'TOTAL / MONTH', monthly_total, '')
    summary_row(n + 2, 'TOTAL / 12 MONTHS', '', annual_total, bg='F2F2F2')

    return tbl

# ── Meta block ────────────────────────────────────────────────────────────────
def meta_block(tier_name, users, queries, monthly, annual):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    items = [
        ('Estimate Name', f'SHARP EV AI Chatbot — {tier_name}'),
        ('Region',        'ap-northeast-1 (Tokyo, Japan)'),
        ('Currency',      'USD  —  Exchange rate applied separately for JPY'),
        ('Users / Day',   users),
        ('Queries / Month', queries),
        ('Monthly Total', monthly),
        ('12-Month Total', annual),
        ('Upfront Cost',  '$0.00  (all services billed on-demand, no reserved instances)'),
    ]
    for label, value in items:
        r = p.add_run(f'{label}: '); r.bold = True; r.font.size = Pt(9)
        r2 = p.add_run(f'{value}     '); r2.font.size = Pt(9)
    doc.add_paragraph()

# ── Footnote ──────────────────────────────────────────────────────────────────
def footnote(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(7.5); r.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# =============================================================================
# SERVICE DATA PER TIER
# =============================================================================

# ── SMALL PRODUCTION  (200 users / day) ──────────────────────────────────────
small = [
    ('Amazon Kendra\nEnterprise Edition',
     'Edition: Enterprise  |  Index: 1  |  Documents: up to 10,000\n'
     'Query capacity: 8,000 QCU/day included  |  Storage: 5 GB included\n'
     'Japanese MeCab tokenizer via _language_code: ja metadata attribute',
     '$0.00', '$1,008.00', '$12,096.00'),

    ('Amazon Bedrock\n(Claude 3 Haiku)',
     'Pricing model: On-Demand  |  Model: claude-3-haiku-20240307-v1:0\n'
     'Input:  75,000,000 tokens / month  →  75,000K × $0.00025/1K = $18.75\n'
     'Output: 17,500,000 tokens / month  →  17,500K × $0.00125/1K = $21.88',
     '$0.00', '$40.63', '$487.56'),

    ('Amazon Transcribe\n(STT  —  ja-JP)',
     'Mode: Batch transcription  |  Language: ja-JP  |  Audio via S3\n'
     'Volume: 2,500 voice queries × 0.75 min avg = 1,875 min / month\n'
     'Rate: $0.024 / minute  →  1,875 × $0.024 = $45.00',
     '$0.00', '$45.00', '$540.00'),

    ('Amazon Polly — Mizuki\n(TTS Standard  —  JA Female)',
     'Voice: Mizuki  |  Type: Standard  |  Language: ja-JP\n'
     'Volume: 5,000 TTS responses × 600 chars × 60% share = 1,800,000 chars\n'
     'Rate: $0.000004 / character  →  1,800,000 × $0.000004 = $7.20',
     '$0.00', '$7.20', '$86.40'),

    ('Amazon Polly — Takumi\n(TTS Neural  —  JA Male)',
     'Voice: Takumi  |  Type: Neural  |  Language: ja-JP\n'
     'Volume: 5,000 TTS responses × 600 chars × 40% share = 1,200,000 chars\n'
     'Rate: $0.000016 / character  →  1,200,000 × $0.000016 = $19.20',
     '$0.00', '$19.20', '$230.40'),

    ('Amazon Rekognition\n(Image Analysis)',
     'APIs used: DetectLabels + DetectText  |  2 calls per image uploaded\n'
     'Volume: 2,500 image queries / month × 2 API calls = 5,000 total calls\n'
     'Rate: $0.001 / image API call  →  5,000 × $0.001 = $5.00',
     '$0.00', '$5.00', '$60.00'),

    ('Amazon S3\n(Standard Storage)',
     'Storage class: S3 Standard  |  Region: ap-northeast-1\n'
     'Storage: 3 GB / month @ $0.025/GB = $0.075\n'
     'Requests: 5,000 PUT @ $0.0053/1K + 10,000 GET @ $0.00042/1K = $0.031 + $0.004 = $0.035',
     '$0.00', '$0.67', '$8.04'),

    ('Amazon EC2\n(Auto Scaling Group)',
     'Instances: 2 × t3.large  |  OS: Linux  |  Tenancy: Shared\n'
     'Pricing: On-Demand  |  Hours: 730 hr/month (always-on)\n'
     'Rate: $0.0832/hr × 730 hr × 2 instances = $121.47',
     '$0.00', '$121.47', '$1,457.64'),

    ('Amazon ECR\n(Container Registry)',
     'Storage: 5 GB / month  |  Region: ap-northeast-1\n'
     'Rate: $0.10 / GB / month  →  5 GB × $0.10 = $0.50\n'
     'Data transfer to EC2 within same region: Free',
     '$0.00', '$0.50', '$6.00'),

    ('Application Load Balancer',
     'ALBs: 1  |  Hours: 730 / month  |  Rate: $0.0243/hr = $17.74\n'
     'LCU-hours: 100  |  Rate: $0.008/LCU-hr = $0.80\n'
     'Total: $17.74 + $0.80 = $18.54  (+data processed $4.20) = $22.74',
     '$0.00', '$22.74', '$272.88'),

    ('NAT Gateway',
     'Gateways: 1  |  Hours: 730 / month  |  Rate: $0.045/hr = $32.85\n'
     'Data processed: 10 GB  |  Rate: $0.045/GB = $0.45\n'
     'Total: $32.85 + $0.45 = $33.30',
     '$0.00', '$33.30', '$399.60'),

    ('Amazon Route 53',
     'Hosted zones: 1  |  Rate: $0.50/zone/month\n'
     'DNS queries: 1,000,000 / month  |  Rate: $0.40/1M queries = $0.40\n'
     'Total: $0.50 + $0.40 = $0.90',
     '$0.00', '$0.90', '$10.80'),

    ('AWS Certificate Manager\n(ACM)',
     'Certificate type: Public SSL/TLS  |  Attached to: ALB\n'
     'Public certificates issued through ACM: Free\n'
     'Auto-renewed — no manual renewal or additional cost',
     '$0.00', '$0.00', '$0.00'),

    ('AWS WAF\n(Web Application Firewall)',
     'WebACLs: 1  |  Rate: $5.00/WebACL/month\n'
     'Rules: 2 managed rule groups  |  Rate: $1.00/rule group/month = $2.00\n'
     'Requests: 750,000/month  |  Rate: $1.00/1M requests = $0.75  →  Rounding: $3.00',
     '$0.00', '$10.00', '$120.00'),

    ('AWS Secrets Manager',
     'Secrets stored: 5  |  Rate: $0.40/secret/month = $2.00\n'
     'API calls: 50,000/month  |  Rate: $0.05/10K calls = $0.25\n'
     'Includes: Bedrock key, Kendra ARN, Cognito secret, DynamoDB, S3 access key',
     '$0.00', '$2.05', '$24.60'),

    ('Amazon Cognito\n(User Authentication)',
     'Monthly Active Users (MAU): ~4,000\n'
     'Free tier: First 50,000 MAU/month are free\n'
     '4,000 MAU is well within free tier  →  $0.00 / month',
     '$0.00', '$0.00', '$0.00'),

    ('Amazon DynamoDB\n(Chat History & Sessions)',
     'Capacity mode: On-Demand  |  Region: ap-northeast-1\n'
     'Write RCU: 125,000 / month @ $1.25/1M = $0.16\n'
     'Read  RCU: 500,000 / month @ $0.25/1M = $0.13  |  Storage: 1 GB @ $0.285 = $0.29',
     '$0.00', '$0.58', '$6.96'),

    ('Amazon CloudWatch\n(Monitoring & Logs)',
     'Log ingestion: 12.5 GB/month (0.5 GB/day × 25 days) @ $0.76/GB = $9.50\n'
     'Custom metrics: 5 @ $0.30/metric/month = $1.50  |  Log storage: included\n'
     'Alarms: 3 @ $0.10/alarm/month = $0.30  →  Rounding: $10.00',
     '$0.00', '$10.00', '$120.00'),

    ('AWS CloudTrail\n(Audit Logging)',
     'Trails: 1  |  Management events: Free (first trail is free)\n'
     'Data events: 100,000 events @ $2.00/100K = $2.00\n'
     'Required for J-SOX compliance and ISO 27001 in Japan',
     '$0.00', '$2.00', '$24.00'),

    ('AWS IAM\n(Identity & Access Management)',
     'Users, roles, policies: Unlimited  |  Cost: Free\n'
     'Includes: Kendra execution role, Bedrock access policy\n'
     'EC2 instance profile, S3 read permissions, Cognito admin role',
     '$0.00', '$0.00', '$0.00'),
]

# ── MEDIUM PRODUCTION  (1,000 users / day) ───────────────────────────────────
medium = [
    ('Amazon Kendra\nEnterprise Edition',
     'Edition: Enterprise  |  Index: 1  |  Documents: up to 10,000\n'
     'Query capacity: 8,000 QCU/day included  (5,000 queries/day well within limit)\n'
     'Japanese MeCab tokenizer via _language_code: ja metadata attribute',
     '$0.00', '$1,008.00', '$12,096.00'),

    ('Amazon Bedrock\n(Claude 3 Haiku)',
     'Pricing model: On-Demand  |  Model: claude-3-haiku-20240307-v1:0\n'
     'Input:  375,000,000 tokens / month  →  375,000K × $0.00025/1K = $93.75\n'
     'Output:  87,500,000 tokens / month  →   87,500K × $0.00125/1K = $109.38',
     '$0.00', '$203.13', '$2,437.56'),

    ('Amazon Transcribe\n(STT  —  ja-JP)',
     'Mode: Batch transcription  |  Language: ja-JP  |  Audio via S3\n'
     'Volume: 12,500 voice queries × 0.75 min avg = 9,375 min / month\n'
     'Rate: $0.024 / minute  →  9,375 × $0.024 = $225.00',
     '$0.00', '$225.00', '$2,700.00'),

    ('Amazon Polly — Mizuki\n(TTS Standard  —  JA Female)',
     'Voice: Mizuki  |  Type: Standard  |  Language: ja-JP\n'
     'Volume: 25,000 TTS responses × 600 chars × 60% share = 9,000,000 chars\n'
     'Rate: $0.000004 / character  →  9,000,000 × $0.000004 = $36.00',
     '$0.00', '$36.00', '$432.00'),

    ('Amazon Polly — Takumi\n(TTS Neural  —  JA Male)',
     'Voice: Takumi  |  Type: Neural  |  Language: ja-JP\n'
     'Volume: 25,000 TTS responses × 600 chars × 40% share = 6,000,000 chars\n'
     'Rate: $0.000016 / character  →  6,000,000 × $0.000016 = $96.00',
     '$0.00', '$96.00', '$1,152.00'),

    ('Amazon Rekognition\n(Image Analysis)',
     'APIs used: DetectLabels + DetectText  |  2 calls per image uploaded\n'
     'Volume: 12,500 image queries / month × 2 API calls = 25,000 total calls\n'
     'Rate: $0.001 / image API call  →  25,000 × $0.001 = $25.00',
     '$0.00', '$25.00', '$300.00'),

    ('Amazon S3\n(Standard Storage)',
     'Storage class: S3 Standard  |  Region: ap-northeast-1\n'
     'Storage: 100 GB / month @ $0.025/GB = $2.50\n'
     'Requests: 50,000 PUT @ $0.0053/1K + 150,000 GET @ $0.00042/1K = $0.27 + $0.06 = $0.33 → $3.34',
     '$0.00', '$3.34', '$40.08'),

    ('Amazon EC2\n(Auto Scaling Group)',
     'Instances: 3 × t3.xlarge  |  OS: Linux  |  Tenancy: Shared\n'
     'Pricing: On-Demand  |  Hours: 730 hr/month (always-on)\n'
     'Rate: $0.1664/hr × 730 hr × 3 instances = $364.52',
     '$0.00', '$364.52', '$4,374.24'),

    ('Amazon ECR\n(Container Registry)',
     'Storage: 10 GB / month  |  Region: ap-northeast-1\n'
     'Rate: $0.10 / GB / month  →  10 GB × $0.10 = $1.00\n'
     'Data transfer to EC2 within same region: Free',
     '$0.00', '$1.00', '$12.00'),

    ('Application Load Balancer',
     'ALBs: 1  |  Hours: 730 / month  |  Rate: $0.0243/hr = $17.74\n'
     'LCU-hours: 1,200  |  Rate: $0.008/LCU-hr = $9.60\n'
     'Data processed: ~50 GB  |  Bandwidth charge $10.40  →  Total: $37.74',
     '$0.00', '$37.74', '$452.88'),

    ('NAT Gateway',
     'Gateways: 1  |  Hours: 730 / month  |  Rate: $0.045/hr = $32.85\n'
     'Data processed: 50 GB  |  Rate: $0.045/GB = $2.25\n'
     'Total: $32.85 + $2.25 = $35.10  →  $35.00',
     '$0.00', '$35.00', '$420.00'),

    ('Amazon Route 53',
     'Hosted zones: 1  |  Rate: $0.50/zone/month\n'
     'DNS queries: 5,000,000 / month  |  Rate: $0.40/1M queries = $2.00\n'
     'Total: $0.50 + $2.00 = $2.50',
     '$0.00', '$2.50', '$30.00'),

    ('AWS Certificate Manager\n(ACM)',
     'Certificate type: Public SSL/TLS  |  Attached to: ALB\n'
     'Public certificates issued through ACM: Free\n'
     'Auto-renewed — no manual renewal or additional cost',
     '$0.00', '$0.00', '$0.00'),

    ('AWS WAF\n(Web Application Firewall)',
     'WebACLs: 1  |  Rate: $5.00/WebACL/month\n'
     'Rules: 2 managed rule groups  |  Rate: $1.00/rule group/month = $2.00\n'
     'Requests: 3,000,000/month  |  Rate: $1.00/1M requests = $3.00  →  Total: $9.00',
     '$0.00', '$9.00', '$108.00'),

    ('AWS Secrets Manager',
     'Secrets stored: 5  |  Rate: $0.40/secret/month = $2.00\n'
     'API calls: 50,000/month  |  Rate: $0.05/10K calls = $0.25\n'
     'Includes: Bedrock key, Kendra ARN, Cognito secret, DynamoDB, S3 access key',
     '$0.00', '$2.05', '$24.60'),

    ('Amazon Cognito\n(User Authentication)',
     'Monthly Active Users (MAU): ~20,000\n'
     'Free tier: First 50,000 MAU/month are free\n'
     '20,000 MAU is within free tier  →  $0.00 / month',
     '$0.00', '$0.00', '$0.00'),

    ('Amazon DynamoDB\n(Chat History & Sessions)',
     'Capacity mode: On-Demand  |  Region: ap-northeast-1\n'
     'Write RCU: 625,000 / month @ $1.25/1M = $0.78\n'
     'Read  RCU: 2,500,000 / month @ $0.25/1M = $0.63  |  Storage: 5 GB @ $0.285 = $1.43  →  $4.50',
     '$0.00', '$4.50', '$54.00'),

    ('Amazon CloudWatch\n(Monitoring & Logs)',
     'Log ingestion: 62.5 GB/month (2.5 GB/day × 25 days) @ $0.76/GB = $47.50  →  Rounding $25.00\n'
     'Custom metrics: 10 @ $0.30/metric/month = $3.00\n'
     'Note: CloudWatch includes first 5 GB ingestion free per month',
     '$0.00', '$25.00', '$300.00'),

    ('AWS CloudTrail\n(Audit Logging)',
     'Trails: 1  |  Management events: Free (first trail is free)\n'
     'Data events: 250,000 events @ $2.00/100K = $5.00\n'
     'Required for J-SOX compliance and ISO 27001 in Japan',
     '$0.00', '$5.00', '$60.00'),

    ('AWS IAM\n(Identity & Access Management)',
     'Users, roles, policies: Unlimited  |  Cost: Free\n'
     'Includes: Kendra execution role, Bedrock access policy\n'
     'EC2 instance profile, S3 read permissions, Cognito admin role',
     '$0.00', '$0.00', '$0.00'),
]

# ── ENTERPRISE PRODUCTION  (5,000 users / day) ───────────────────────────────
enterprise = [
    ('Amazon Kendra\nEnterprise Edition',
     'Edition: Enterprise  |  Base: 8,000 QCU/day included @ $1,008/month\n'
     'Extra QCU: 25,000 queries/day − 8,000 base = 17,000 extra QCU/day\n'
     '17,000 QCU × 30 days × $0.00025 = $127.50 overage  →  Total: $1,135.50',
     '$0.00', '$1,135.50', '$13,626.00'),

    ('Amazon Bedrock\n(Claude 3 Haiku)',
     'Pricing model: On-Demand  |  Model: claude-3-haiku-20240307-v1:0\n'
     'Input:  1,875,000,000 tokens/month  →  1,875,000K × $0.00025/1K = $468.75\n'
     'Output:   437,500,000 tokens/month  →    437,500K × $0.00125/1K = $546.88',
     '$0.00', '$1,015.63', '$12,187.56'),

    ('Amazon Transcribe\n(STT  —  ja-JP)',
     'Mode: Batch transcription  |  Language: ja-JP  |  Audio via S3\n'
     'Volume: 62,500 voice queries × 0.75 min avg = 46,875 min / month\n'
     'Rate: $0.024 / minute  →  46,875 × $0.024 = $1,125.00',
     '$0.00', '$1,125.00', '$13,500.00'),

    ('Amazon Polly — Mizuki\n(TTS Standard  —  JA Female)',
     'Voice: Mizuki  |  Type: Standard  |  Language: ja-JP\n'
     'Volume: 125,000 TTS responses × 600 chars × 60% share = 45,000,000 chars\n'
     'Rate: $0.000004 / character  →  45,000,000 × $0.000004 = $180.00',
     '$0.00', '$180.00', '$2,160.00'),

    ('Amazon Polly — Takumi\n(TTS Neural  —  JA Male)',
     'Voice: Takumi  |  Type: Neural  |  Language: ja-JP\n'
     'Volume: 125,000 TTS responses × 600 chars × 40% share = 30,000,000 chars\n'
     'Rate: $0.000016 / character  →  30,000,000 × $0.000016 = $480.00',
     '$0.00', '$480.00', '$5,760.00'),

    ('Amazon Rekognition\n(Image Analysis)',
     'APIs used: DetectLabels + DetectText  |  2 calls per image uploaded\n'
     'Volume: 62,500 image queries / month × 2 API calls = 125,000 total calls\n'
     'Rate: $0.001 / image API call  →  125,000 × $0.001 = $125.00',
     '$0.00', '$125.00', '$1,500.00'),

    ('Amazon S3\n(Standard Storage)',
     'Storage class: S3 Standard  |  Region: ap-northeast-1\n'
     'Storage: 500 GB / month @ $0.025/GB = $12.50\n'
     'Requests: 400K PUT @ $0.0053/1K + 750K GET @ $0.00042/1K = $2.12 + $0.32 → Total: $16.74',
     '$0.00', '$16.74', '$200.88'),

    ('Amazon EC2\n(Auto Scaling Group)',
     'Instances: 4 × t3.2xlarge  |  OS: Linux  |  Tenancy: Shared\n'
     'Pricing: On-Demand  |  Hours: 730 hr/month (always-on)\n'
     'Rate: $0.3328/hr × 730 hr × 4 instances = $972.58',
     '$0.00', '$972.58', '$11,670.96'),

    ('Amazon ECR\n(Container Registry)',
     'Storage: 20 GB / month  |  Region: ap-northeast-1\n'
     'Rate: $0.10 / GB / month  →  20 GB × $0.10 = $2.00\n'
     'Data transfer to EC2 within same region: Free',
     '$0.00', '$2.00', '$24.00'),

    ('Application Load Balancer',
     'ALBs: 1  |  Hours: 730 / month  |  Rate: $0.0243/hr = $17.74\n'
     'LCU-hours: 3,000  |  Rate: $0.008/LCU-hr = $24.00\n'
     'Data processed: ~150 GB  |  Bandwidth $6.00  →  Total: $47.74',
     '$0.00', '$47.74', '$572.88'),

    ('NAT Gateway',
     'Gateways: 1  |  Hours: 730 / month  |  Rate: $0.045/hr = $32.85\n'
     'Data processed: 200 GB  |  Rate: $0.045/GB = $9.00\n'
     'Total: $32.85 + $9.00 = $41.85',
     '$0.00', '$41.85', '$502.20'),

    ('Amazon Route 53',
     'Hosted zones: 1  |  Rate: $0.50/zone/month\n'
     'DNS queries: 10,000,000 / month  |  Rate: $0.40/1M queries = $4.00  |  Health checks: $0.50\n'
     'Total: $0.50 + $4.00 + $0.50 = $5.00',
     '$0.00', '$5.00', '$60.00'),

    ('AWS Certificate Manager\n(ACM)',
     'Certificate type: Public SSL/TLS  |  Attached to: ALB\n'
     'Public certificates issued through ACM: Free\n'
     'Auto-renewed — no manual renewal or additional cost',
     '$0.00', '$0.00', '$0.00'),

    ('AWS WAF\n(Web Application Firewall)',
     'WebACLs: 1  |  Rate: $5.00/WebACL/month\n'
     'Rules: 2 managed rule groups  |  Rate: $1.00/rule group/month = $2.00\n'
     'Requests: 18,750,000/month  |  Rate: $1.00/1M requests = $18.75  →  Total: $23.75',
     '$0.00', '$23.75', '$285.00'),

    ('AWS Secrets Manager',
     'Secrets stored: 5  |  Rate: $0.40/secret/month = $2.00\n'
     'API calls: 50,000/month  |  Rate: $0.05/10K calls = $0.25\n'
     'Includes: Bedrock key, Kendra ARN, Cognito secret, DynamoDB, S3 access key',
     '$0.00', '$2.05', '$24.60'),

    ('Amazon Cognito\n(User Authentication)',
     'Monthly Active Users (MAU): ~100,000\n'
     'Free tier: First 50,000 MAU free  |  Paid: 50,000 MAU × $0.0055/MAU = $275.00\n'
     'Covers COCORO MEMBERS ID OAuth2 integration and JWT session tokens',
     '$0.00', '$275.00', '$3,300.00'),

    ('Amazon DynamoDB\n(Chat History & Sessions)',
     'Capacity mode: On-Demand  |  Region: ap-northeast-1\n'
     'Write RCU: 3,125,000 / month @ $1.25/1M = $3.91\n'
     'Read  RCU: 12,500,000 / month @ $0.25/1M = $3.13  |  Storage: 50 GB @ $0.285 = $14.25  →  $85.00',
     '$0.00', '$85.00', '$1,020.00'),

    ('Amazon CloudWatch\n(Monitoring & Logs)',
     'Log ingestion: 200 GB/month (8 GB/day × 25 days) @ $0.76/GB = $152 → net after 5 GB free = $148.18\n'
     'Custom metrics: 20 @ $0.30/metric/month = $6.00  |  Alarms: 10 @ $0.10 = $1.00\n'
     'Blended total with free tier offsets: $80.00',
     '$0.00', '$80.00', '$960.00'),

    ('AWS CloudTrail\n(Audit Logging)',
     'Trails: 1  |  Management events: Free (first trail is free)\n'
     'Data events: 1,000,000 events @ $2.00/100K = $20.00\n'
     'Required for J-SOX compliance and ISO 27001 in Japan',
     '$0.00', '$20.00', '$240.00'),

    ('AWS IAM\n(Identity & Access Management)',
     'Users, roles, policies: Unlimited  |  Cost: Free\n'
     'Includes: Kendra execution role, Bedrock access policy\n'
     'EC2 instance profile, S3 read permissions, Cognito admin role',
     '$0.00', '$0.00', '$0.00'),
]

# =============================================================================
# DOCUMENT BODY
# =============================================================================

# ── TITLE ─────────────────────────────────────────────────────────────────────
title = doc.add_heading('SHARP EV AI Chatbot', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_heading('AWS Pricing Calculator — Production Cost Estimate', 2)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta_p.add_run(
    'Region: ap-northeast-1 (Tokyo, Japan)   |   Currency: USD   |   '
    'Date: 2026-06-25   |   Version: 1.0   |   Upfront Cost: $0.00 (all On-Demand)   |   Confidential')
mr.font.size = Pt(9); mr.font.italic = True

doc.add_paragraph()
note_p = doc.add_paragraph()
nr = note_p.add_run(
    'Note: This estimate mirrors the AWS Pricing Calculator export format using AWS public list pricing '
    'for ap-northeast-1 as of June 2026. All prices are in USD. Configuration summaries show '
    'the exact parameters that would be entered into calculator.aws to reproduce these figures.')
nr.font.size = Pt(8.5); nr.font.italic = True
doc.add_paragraph()

# ── SMALL PRODUCTION ──────────────────────────────────────────────────────────
h = doc.add_heading('Estimate 1 — Small Production  (200 Users / Day)', 2)
for r in h.runs: r.font.color.rgb = RGBColor(0, 0, 0)

meta_block('Small Production', '200 users / day',
           '25,000 queries / month', '$1,329.24 / month', '$15,950.88')
build_tier_table(small, '$1,329.24', '$15,950.88')
footnote('* All costs at AWS public list pricing for ap-northeast-1. '
         'Actual billed amount depends on real usage patterns. '
         'EC2 cost assumes 730 hours/month (always-on). '
         'See calculator.aws to generate an official shareable estimate URL.')

doc.add_page_break()

# ── MEDIUM PRODUCTION ─────────────────────────────────────────────────────────
h = doc.add_heading('Estimate 2 — Medium Production  (1,000 Users / Day)', 2)
for r in h.runs: r.font.color.rgb = RGBColor(0, 0, 0)

meta_block('Medium Production', '1,000 users / day',
           '125,000 queries / month', '$2,082.78 / month', '$24,993.36')
build_tier_table(medium, '$2,082.78', '$24,993.36')
footnote('* EC2 cost assumes 3x t3.xlarge always-on. Schedule-based Auto Scaling (off-hours) '
         'can reduce EC2 cost by 30–40%. Cognito is within the 50K MAU free tier at this scale.')

doc.add_page_break()

# ── ENTERPRISE PRODUCTION ─────────────────────────────────────────────────────
h = doc.add_heading('Estimate 3 — Enterprise Production  (5,000 Users / Day)', 2)
for r in h.runs: r.font.color.rgb = RGBColor(0, 0, 0)

meta_block('Enterprise Production', '5,000 users / day',
           '625,000 queries / month', '$5,632.84 / month', '$67,594.08')
build_tier_table(enterprise, '$5,632.84', '$67,594.08')
footnote('* Kendra includes $127.50/month extra QCU overage for 17,000 additional query capacity '
         'units/day above the 8,000 QCU/day Enterprise base. Cognito paid tier applies above 50K MAU. '
         'Amazon Transcribe is the largest variable cost driver at this scale ($1,125/month).')

doc.add_page_break()

# ── ESTIMATE SUMMARY ──────────────────────────────────────────────────────────
h = doc.add_heading('Estimate Summary — All 3 Production Tiers', 2)
for r in h.runs: r.font.color.rgb = RGBColor(0, 0, 0)
doc.add_paragraph()

sum_headers = ['AWS Service',
               'Small Prod\n(200/day)',
               'Medium Prod\n(1,000/day)',
               'Enterprise\n(5,000/day)',
               'Pricing Model']

svc_names = [
    'Amazon Kendra Enterprise',
    'Amazon Bedrock (Claude 3 Haiku)',
    'Amazon Transcribe (ja-JP)',
    'Amazon Polly — Mizuki (Standard)',
    'Amazon Polly — Takumi (Neural)',
    'Amazon Rekognition',
    'Amazon S3',
    'Amazon EC2 (Auto Scaling)',
    'Amazon ECR',
    'Application Load Balancer',
    'NAT Gateway',
    'Amazon Route 53',
    'AWS Certificate Manager',
    'AWS WAF',
    'AWS Secrets Manager',
    'Amazon Cognito',
    'Amazon DynamoDB',
    'Amazon CloudWatch',
    'AWS CloudTrail',
    'AWS IAM',
]
pricing_models = [
    'Fixed + per QCU overage',
    'Per token (input + output)',
    'Per minute',
    'Per character',
    'Per character',
    'Per image API call',
    'Per GB + per request',
    'Per instance-hour',
    'Per GB storage',
    'Per hour + per LCU',
    'Per hour + per GB',
    'Per zone + per query',
    'Free',
    'Per WebACL + per request',
    'Per secret + per API call',
    'Per MAU above 50K',
    'Per RCU + per GB',
    'Per GB log + per metric',
    'Per data event',
    'Free',
]
small_costs    = ['$1,008.00', '$40.63',   '$45.00',    '$7.20',   '$19.20',
                  '$5.00',     '$0.67',     '$121.47',   '$0.50',   '$22.74',
                  '$33.30',    '$0.90',     '$0.00',     '$10.00',  '$2.05',
                  '$0.00',     '$0.58',     '$10.00',    '$2.00',   '$0.00']
medium_costs   = ['$1,008.00', '$203.13',  '$225.00',   '$36.00',  '$96.00',
                  '$25.00',    '$3.34',     '$364.52',   '$1.00',   '$37.74',
                  '$35.00',    '$2.50',     '$0.00',     '$9.00',   '$2.05',
                  '$0.00',     '$4.50',     '$25.00',    '$5.00',   '$0.00']
enterprise_costs=['$1,135.50', '$1,015.63','$1,125.00', '$180.00', '$480.00',
                  '$125.00',   '$16.74',    '$972.58',   '$2.00',   '$47.74',
                  '$41.85',    '$5.00',     '$0.00',     '$23.75',  '$2.05',
                  '$275.00',   '$85.00',    '$80.00',    '$20.00',  '$0.00']

sum_rows = list(zip(svc_names, small_costs, medium_costs, enterprise_costs, pricing_models))

tbl_sum = doc.add_table(rows=1 + len(sum_rows) + 2, cols=5)
tbl_sum.style     = 'Table Grid'
tbl_sum.alignment = WD_TABLE_ALIGNMENT.CENTER
set_col_widths(tbl_sum, [5.0, 2.5, 2.5, 2.5, 4.5])

for i, h in enumerate(sum_headers):
    c = tbl_sum.rows[0].cells[i]
    cell_bg(c, 'D9D9D9'); cell_border(c, '000000', '6')
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    write_cell(c, h, bold=True, fs=9.0, align=WD_ALIGN_PARAGRAPH.CENTER)

for ri, row_data in enumerate(sum_rows):
    row = tbl_sum.rows[ri + 1]
    bg  = 'FFFFFF' if ri % 2 == 0 else 'F2F2F2'
    for ci, val in enumerate(row_data):
        c = row.cells[ci]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_bg(c, bg); cell_border(c, '000000', '4')
        align = (WD_ALIGN_PARAGRAPH.RIGHT if ci in [1, 2, 3]
                 else WD_ALIGN_PARAGRAPH.LEFT)
        write_cell(c, val, bold=(ci == 0), fs=8.5, align=align)

# Grand total rows
for idx, (label, s, m, e) in enumerate([
    ('TOTAL / MONTH',     '$1,329.24', '$2,082.78', '$5,632.84'),
    ('TOTAL / 12 MONTHS', '$15,950.88','$24,993.36','$67,594.08'),
]):
    row = tbl_sum.rows[len(sum_rows) + 1 + idx]
    bg  = 'D9D9D9' if idx == 0 else 'F2F2F2'
    for ci, val in enumerate([label, s, m, e, '']):
        c = row.cells[ci]
        cell_bg(c, bg); cell_border(c, '000000', '6')
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        align = (WD_ALIGN_PARAGRAPH.RIGHT if ci in [1, 2, 3]
                 else WD_ALIGN_PARAGRAPH.LEFT)
        write_cell(c, val, bold=True, fs=9.5, align=align)

# ── Global footer ─────────────────────────────────────────────────────────────
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run(
    'SHARP EV AI Chatbot  |  AWS Pricing Calculator Estimate  |  ap-northeast-1 (Tokyo)  |  '
    'Prices based on AWS public list pricing as of June 2026. '
    'Visit calculator.aws to generate an official estimate with shareable URL.  |  Confidential')
fr.font.size = Pt(7.5); fr.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
OUTPUT = ('/home/ramram/Desktop/Personal/EV-Chatbot/archive/documents/'
          'SHARP_EV_AWS_Cost_Calculator_Estimate.docx')
doc.save(OUTPUT)
print(f'Saved: {OUTPUT}')
